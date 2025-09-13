from docx import Document

# Very minimal template renderer: replaces {{key}} in all paragraphs and table cells

def render_docx_template(template_path: str, context: dict[str, str], output_path: str) -> None:
    doc = Document(template_path)
    # Replace in paragraphs
    for p in doc.paragraphs:
        for k, v in context.items():
            if f"{{{{{k}}}}}" in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    inline[i].text = inline[i].text.replace(f"{{{{{k}}}}}", v)
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for k, v in context.items():
                    if f"{{{{{k}}}}}" in cell.text:
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.text = r.text.replace(f"{{{{{k}}}}}", v)
    doc.save(output_path)
