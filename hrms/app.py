from PySide6.QtWidgets import (
    QApplication, QWidget, QLineEdit, QPushButton, QVBoxLayout, QLabel,
    QListWidget, QMessageBox, QHBoxLayout, QComboBox, QDialog, QFormLayout
)
from PySide6.QtCore import Qt, QTimer
import sys
from pathlib import Path

from docx import Document

from .db import SessionLocal
from .models import Person, User
from .init_db import init_db
from .seed import seed_basic_data
from .security import verify_password
from .salary import list_due_in_window, export_due_to_excel
from .audit import log_action
from .scheduler import NOTIFY_QUEUE
from .templates import render_docx_template
from .ui_forms import prompt_context
from .person_detail import PersonDetailDialog


class LoginWindow(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("HRMS - Đăng nhập")
        layout = QVBoxLayout()
        self.user = QLineEdit()
        self.user.setPlaceholderText("Username")
        self.pw = QLineEdit()
        self.pw.setPlaceholderText("Password")
        self.pw.setEchoMode(QLineEdit.Password)
        self.btn = QPushButton("Login")
        self.btn.clicked.connect(self.do_login)
        layout.addWidget(self.user)
        layout.addWidget(self.pw)
        layout.addWidget(self.btn)
        self.setLayout(layout)

    def do_login(self):
        username = self.user.text().strip()
        password = self.pw.text()
        if not username or not password:
            QMessageBox.warning(self, "Thiếu thông tin", "Vui lòng nhập username và password")
            return
        db = SessionLocal()
        try:
            user = db.query(User).filter(User.username == username).first()
            if not user or not verify_password(password, user.password_hash):
                QMessageBox.critical(self, "Đăng nhập thất bại", "Sai thông tin đăng nhập")
                return
            user_info = {"id": user.id, "username": user.username, "role": user.role, "unit_id": user.unit_id}
        finally:
            db.close()
        self.main = MainWindow(current_user=user_info)
        self.main.show()
        self.close()


class MainWindow(QWidget):
    def __init__(self, current_user: dict):
        super().__init__()
        self.current_user = current_user
        self.setWindowTitle(f"HRMS - Tra cứu nhân sự ({current_user.get('username')})")
        layout = QVBoxLayout()

        # Hàng filter
        filter_bar = QHBoxLayout()
        self.filter_unit = QComboBox()
        self.filter_unit.addItem("-- Đơn vị --", None)
        self.filter_position = QComboBox()
        self.filter_position.addItem("-- Chức vụ --", None)
        self.filter_rank = QComboBox(); self.filter_rank.addItem("-- Ngạch --", None)
        self.filter_step = QComboBox(); self.filter_step.addItem("-- Bậc --", None)
        self.filter_status = QComboBox()
        self.filter_status.addItems(["-- Trạng thái --", "Đang công tác", "Nghỉ thai sản", "Đi học", "Thôi việc"])  # có thể mở rộng
        self.btn_filter_refresh = QPushButton("Lọc")
        self.btn_filter_refresh.clicked.connect(lambda: self.on_search(self.search.text()))
        self.btn_manage_users = QPushButton("Quản lý người dùng")
        self.btn_manage_users.clicked.connect(self.manage_users)
        self.btn_settings = QPushButton("Cấu hình")
        self.btn_settings.clicked.connect(self.open_settings)
        self.btn_email_history = QPushButton("Lịch sử Email")
        self.btn_email_history.clicked.connect(self.open_email_history)
        filter_bar.addWidget(self.filter_unit)
        filter_bar.addWidget(self.filter_position)
        filter_bar.addWidget(self.filter_rank)
        filter_bar.addWidget(self.filter_step)
        filter_bar.addWidget(self.filter_status)
        filter_bar.addWidget(self.btn_filter_refresh)
        filter_bar.addWidget(self.btn_manage_users)
        filter_bar.addWidget(self.btn_settings)

        self.search = QLineEdit()
        self.search.setPlaceholderText("Nhập tên cần tìm...")
        self.search.textChanged.connect(self.on_search)
        self.list = QListWidget()
        # Phân trang
        pager_layout = QHBoxLayout()
        self.page_size_box = QComboBox(); self.page_size_box.addItems(["50","100","200"]); self.page_size_box.setCurrentText("50")
        self.btn_prev = QPushButton("◀ Trước")
        self.btn_next = QPushButton("Sau ▶")
        self.page_label = QLabel("Trang 1/1")
        pager_layout.addWidget(QLabel("Kích thước trang"))
        pager_layout.addWidget(self.page_size_box)
        pager_layout.addWidget(self.btn_prev)
        pager_layout.addWidget(self.btn_next)
        pager_layout.addWidget(self.page_label)
        self.current_page = 0
        self.total_count = 0
        self.btn_prev.clicked.connect(self.go_prev_page)
        self.btn_next.clicked.connect(self.go_next_page)
        self.page_size_box.currentTextChanged.connect(lambda _: self.reset_to_first_page())
        btn_layout = QHBoxLayout()
        self.btn_export = QPushButton("Xuất trích ngang")
        self.btn_export.clicked.connect(self.export_selected)
        self.btn_detail = QPushButton("Chi tiết")
        self.btn_detail.clicked.connect(self.open_detail)
        self.btn_due = QPushButton("Nâng lương quý này")
        self.btn_due.clicked.connect(self.show_due)
        self.btn_email_salary_due = QPushButton("Gửi email nâng lương (quý)")
        self.btn_email_salary_due.clicked.connect(self.send_quarter_salary_email)
        self.btn_email_salary_due_units = QPushButton("Gửi nâng lương theo đơn vị")
        self.btn_email_salary_due_units.clicked.connect(self.send_quarter_salary_email_by_unit)
        self.btn_appointment = QPushButton("Kiểm tra bổ nhiệm")
        self.btn_appointment.clicked.connect(self.check_appointment)
        self.btn_work = QPushButton("Thêm quá trình công tác")
        self.btn_work.clicked.connect(self.add_work_process)
        self.btn_export_work = QPushButton("Xuất quá trình công tác")
        self.btn_export_work.clicked.connect(self.export_work_process)
        self.btn_report = QPushButton("Báo cáo nhanh")
        self.btn_report.clicked.connect(self.quick_report)
        self.btn_email_report = QPushButton("Gửi báo cáo nhanh (email)")
        self.btn_email_report.clicked.connect(self.send_quick_report_email)
        self.btn_letter_salary_cover = QPushButton("CV rà soát nâng lương")
        self.btn_letter_salary_cover.clicked.connect(self.export_salary_cover)
        self.btn_letter_retirement = QPushButton("Văn bản nghỉ hưu")
        self.btn_letter_retirement.clicked.connect(self.export_retirement_letters)
        self.btn_letter_salary_person = QPushButton("Văn bản nâng lương")
        self.btn_letter_salary_person.clicked.connect(self.export_salary_letters)
        self.btn_pdf_toggle = QComboBox(); self.btn_pdf_toggle.addItems(["DOCX","PDF"])  # chế độ xuất
        # Template Excel chooser
        self.excel_template_box = QComboBox(); self.load_excel_templates(); self.load_default_excel_template()
        # Combobox loại export để lưu mặc định theo loại
        self.excel_template_type_box = QComboBox(); self._init_excel_template_types()
        self.btn_save_tpl = QPushButton("Lưu mặc định"); self.btn_save_tpl.clicked.connect(self.save_default_excel_template)
        self.btn_import = QPushButton("Import Excel nhân sự")
        self.btn_import.clicked.connect(self.import_excel)
        self.btn_ins_add = QPushButton("Thêm sự kiện BHXH")
        self.btn_ins_add.clicked.connect(self.add_insurance)
        self.btn_ins_export = QPushButton("Xuất Excel BHXH")
        self.btn_ins_export.clicked.connect(self.export_insurance)
        # Gửi BHXH theo đơn vị (UI)
        self.btn_email_insurance_units = QPushButton("Gửi BHXH theo đơn vị")
        self.btn_email_insurance_units.clicked.connect(self.send_insurance_email_by_unit)
        self.btn_contract_add = QPushButton("Thêm HĐ")
        self.btn_contract_add.clicked.connect(self.add_contract)
        self.btn_contract_export = QPushButton("Xuất HĐ")
        self.btn_contract_export.clicked.connect(self.export_contract)
        # Gửi HĐ sắp hết hạn theo đơn vị (UI)
        self.btn_email_contracts_units = QPushButton("Gửi HĐ sắp hết hạn theo đơn vị")
        self.btn_email_contracts_units.clicked.connect(self.send_contracts_expiring_email_by_unit)
        self.btn_salary_export_filtered = QPushButton("Xuất lương (lọc)")
        self.btn_salary_export_filtered.clicked.connect(self.export_salary_histories_filtered)
        btn_layout.addWidget(self.btn_export)
        btn_layout.addWidget(self.btn_detail)
        btn_layout.addWidget(self.btn_due)
        btn_layout.addWidget(self.btn_email_salary_due)
        btn_layout.addWidget(self.btn_email_salary_due_units)
        btn_layout.addWidget(self.btn_appointment)
        btn_layout.addWidget(self.btn_work)
        btn_layout.addWidget(self.btn_export_work)
        btn_layout.addWidget(self.btn_contract_add)
        btn_layout.addWidget(self.btn_contract_export)
        btn_layout.addWidget(self.btn_email_contracts_units)
        btn_layout.addWidget(self.btn_report)
        btn_layout.addWidget(self.btn_email_report)
        btn_layout.addWidget(self.btn_letter_salary_cover)
        btn_layout.addWidget(self.btn_letter_retirement)
        btn_layout.addWidget(self.btn_letter_salary_person)
        btn_layout.addWidget(self.btn_pdf_toggle)
        btn_layout.addWidget(self.excel_template_box)
        btn_layout.addWidget(self.excel_template_type_box)
        btn_layout.addWidget(self.btn_save_tpl)
        btn_layout.addWidget(self.btn_import)
        btn_layout.addWidget(self.btn_email_history)
        btn_layout.addWidget(self.btn_ins_add)
        btn_layout.addWidget(self.btn_ins_export)
        btn_layout.addWidget(self.btn_email_insurance_units)
        btn_layout.addWidget(self.btn_salary_export_filtered)
        layout.addWidget(QLabel("Tra cứu nhân sự"))
        layout.addWidget(self.search)
        layout.addWidget(self.list)
        layout.addLayout(pager_layout)
        layout.addLayout(btn_layout)
        self.setLayout(layout)
        self.setLayout(layout)

        # RBAC: áp dụng quyền theo role
        self.apply_role_permissions()

        # Tải danh mục filter
        self.load_filters()

        # Thiết lập timer để lấy thông báo từ scheduler và popup
        self.timer = QTimer(self)
        self.timer.setInterval(2000)
        self.timer.timeout.connect(self.drain_notifications)
        self.timer.start()

        self.refresh()

    def refresh(self):
        self.on_search("")

    def reset_to_first_page(self):
        self.current_page = 0
        self.on_search(self.search.text())

    def go_prev_page(self):
        if self.current_page > 0:
            self.current_page -= 1
            self.on_search(self.search.text())

    def go_next_page(self):
        page_size = int(self.page_size_box.currentText()) if self.page_size_box.currentText().isdigit() else 50
        max_page = max(0, (self.total_count - 1) // page_size)
        if self.current_page < max_page:
            self.current_page += 1
            self.on_search(self.search.text())

    def on_search(self, text: str):
        db = SessionLocal()
        try:
            q = db.query(Person)
            # Lọc theo tên
            if text:
                like = f"%{text}%"
                q = q.filter(Person.full_name.ilike(like))
            # Lọc theo đơn vị/chức vụ/trạng thái
            unit_id = self.filter_unit.currentData()
            if unit_id:
                q = q.filter(Person.unit_id == unit_id)
            pos_id = self.filter_position.currentData()
            if pos_id:
                q = q.filter(Person.position_id == pos_id)
            # Lọc theo ngạch/bậc dựa trên lịch sử lương mới nhất (tối ưu hóa ở DB)
            rank_id = self.filter_rank.currentData()
            step_val = self.filter_step.currentData()
            if rank_id or step_val:
                from .models import SalaryHistory
                from sqlalchemy import func
                sub = db.query(
                    SalaryHistory.person_id.label('pid'),
                    func.max(SalaryHistory.effective_date).label('max_date')
                ).group_by(SalaryHistory.person_id).subquery()
                latest = db.query(
                    SalaryHistory.person_id.label('pid'),
                    SalaryHistory.rank_id.label('rank_id'),
                    SalaryHistory.step.label('step')
                ).join(
                    sub,
                    (SalaryHistory.person_id == sub.c.pid) & (SalaryHistory.effective_date == sub.c.max_date)
                ).subquery()
                q = q.join(latest, Person.id == latest.c.pid)
                if rank_id:
                    q = q.filter(latest.c.rank_id == rank_id)
                if step_val:
                    q = q.filter(latest.c.step == step_val)
            st = self.filter_status.currentText()
            if st and not st.startswith("--"):
                q = q.filter(Person.status.ilike(f"%{st}%"))
            # Đếm tổng số kết quả trước khi phân trang
            self.total_count = q.count()
            page_size = int(self.page_size_box.currentText()) if self.page_size_box.currentText().isdigit() else 50
            offset = self.current_page * page_size
            # Nếu offset vượt quá tổng, đưa về trang cuối
            if offset >= max(0, self.total_count):
                self.current_page = max(0, (self.total_count - 1) // page_size) if self.total_count > 0 else 0
                offset = self.current_page * page_size
            people = q.order_by(Person.full_name).limit(page_size).offset(offset).all()
            # Cập nhật danh sách
            self.list.clear()
            for p in people:
                self.list.addItem(f"{p.full_name} - {p.code}")
            # Cập nhật pager label và nút
            total_pages = max(1, (self.total_count - 1) // page_size + 1) if self.total_count > 0 else 1
            self.page_label.setText(f"Trang {self.current_page+1}/{total_pages} ({self.total_count} kết quả)")
            self.btn_prev.setEnabled(self.current_page > 0)
            self.btn_next.setEnabled(self.current_page < total_pages - 1)
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", str(e))
        finally:
            db.close()

    def get_selected_excel_template(self) -> str | None:
        # Trả về tên file template (không gồm đường dẫn) hoặc None
        data = self.excel_template_box.currentData()
        return data if isinstance(data, str) or data is None else None

    def _init_excel_template_types(self):
        # Khởi tạo danh sách loại export
        self.excel_template_type_box.clear()
        self.excel_template_type_box.addItem("(Chung)", "GLOBAL")
        self.excel_template_type_box.addItem("Nâng lương đến hạn", "salary_due")
        self.excel_template_type_box.addItem("Lịch sử lương (cá nhân)", "salary_history")
        self.excel_template_type_box.addItem("Lịch sử lương (lọc)", "salary_histories")
        self.excel_template_type_box.addItem("Hợp đồng", "contracts")
        self.excel_template_type_box.addItem("BHXH", "bhxh")

    def get_template_for(self, export_type: str) -> str | None:
        # Ưu tiên mặc định theo loại -> mặc định chung -> lựa chọn hiện tại
        try:
            from .settings_service import get_setting
            username = (self.current_user.get('username') or '').strip()
            if username:
                key_type = f"DEFAULT_XLSX_TEMPLATE:{username}:{export_type}"
                val = get_setting(key_type, None)
                if val:
                    return val
                key_global = f"DEFAULT_XLSX_TEMPLATE:{username}:GLOBAL"
                val2 = get_setting(key_global, None)
                if val2:
                    return val2
        except Exception:
            pass
        return self.get_selected_excel_template()

    def load_excel_templates(self):
        # Nạp danh sách template Excel từ templates/xlsx
        try:
            self.excel_template_box.clear()
            self.excel_template_box.addItem("(Mặc định)", None)
            tpl_dir = Path('templates') / 'xlsx'
            if tpl_dir.exists():
                for p in sorted([x for x in tpl_dir.iterdir() if x.suffix.lower() == '.xlsx']):
                    self.excel_template_box.addItem(p.name, p.name)
        except Exception:
            # Bỏ qua lỗi nạp template
            self.excel_template_box.clear()
            self.excel_template_box.addItem("(Mặc định)", None)

    def load_default_excel_template(self) -> None:
        # Nạp template mặc định (chung) đã lưu theo user hiện tại
        try:
            from .settings_service import get_setting
            username = (self.current_user.get('username') or '').strip()
            if not username:
                return
            key = f"DEFAULT_XLSX_TEMPLATE:{username}:GLOBAL"
            val = get_setting(key, None)
            if not val:
                return
            # Tìm item có data == val
            for i in range(self.excel_template_box.count()):
                if self.excel_template_box.itemData(i) == val:
                    self.excel_template_box.setCurrentIndex(i)
                    break
        except Exception:
            pass

    def save_default_excel_template(self) -> None:
        # Lưu template hiện chọn làm mặc định cho user (theo loại export hoặc GLOBAL)
        try:
            from .settings_service import set_setting
            tpl = self.get_selected_excel_template()
            username = (self.current_user.get('username') or '').strip()
            if not username:
                QMessageBox.warning(self, "Lỗi", "Không xác định được người dùng")
                return
            etype = self.excel_template_type_box.currentData() or 'GLOBAL'
            key = f"DEFAULT_XLSX_TEMPLATE:{username}:{etype}"
            set_setting(key, tpl or '')
            QMessageBox.information(self, "Đã lưu", "Đã lưu template Excel mặc định cho loại export")
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", str(e))

    def load_filters(self):
        db = SessionLocal()
        try:
            self.filter_unit.blockSignals(True)
            self.filter_position.blockSignals(True)
            # Units
            from .models import Unit, Position
            self.filter_unit.clear()
            self.filter_unit.addItem("-- Đơn vị --", None)
            role = (self.current_user.get('role') or '').lower()
            only_unit_id = self.current_user.get('unit_id') if role not in ('admin','hr') else None
            units_q = db.query(Unit).order_by(Unit.name)
            if only_unit_id:
                units_q = units_q.filter(Unit.id == only_unit_id)
            for u in units_q.all():
                self.filter_unit.addItem(u.name, u.id)
            # Positions
            self.filter_position.clear()
            self.filter_position.addItem("-- Chức vụ --", None)
            for p in db.query(Position).order_by(Position.name).all():
                self.filter_position.addItem(p.name, p.id)
            # Ranks
            from .models import SalaryRank, SalaryStep
            self.filter_rank.clear(); self.filter_rank.addItem("-- Ngạch --", None)
            for r in db.query(SalaryRank).order_by(SalaryRank.code).all():
                self.filter_rank.addItem(f"{r.code}-{r.name}", r.id)
            self.filter_step.clear(); self.filter_step.addItem("-- Bậc --", None)
            for s in range(1, 15):
                self.filter_step.addItem(str(s), s)
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", str(e))
        finally:
            self.filter_unit.blockSignals(False)
            self.filter_position.blockSignals(False)
            db.close()

    def apply_role_permissions(self):
        role = (self.current_user.get('role') or '').lower()
        is_admin = role in ('admin', 'hr')
        # Người dùng thường: chỉ xem và xuất trích ngang, xem nâng lương/báo cáo
        self.btn_import.setEnabled(is_admin)
        self.btn_work.setEnabled(is_admin)
        self.btn_detail.setEnabled(True)
        self.btn_export_work.setEnabled(is_admin)
        self.btn_ins_add.setEnabled(is_admin)
        self.btn_ins_export.setEnabled(is_admin)
        self.btn_manage_users.setEnabled(is_admin)
        self.btn_settings.setEnabled(is_admin)
        self.btn_email_history.setEnabled(role in ('admin','hr'))
        # Gửi email nâng lương (quý) chỉ cho admin/hr
        self.btn_email_salary_due.setEnabled(is_admin)
        self.btn_email_salary_due_units.setEnabled(is_admin)
        # Gửi BHXH/HĐ theo đơn vị chỉ cho admin/hr
        self.btn_email_insurance_units.setEnabled(is_admin)
        self.btn_email_contracts_units.setEnabled(is_admin)
        # Nút email nghỉ hưu: chỉ admin/hr
        if not hasattr(self, 'btn_email_retirement'):
            self.btn_email_retirement = QPushButton("Gửi email nghỉ hưu (6/3 tháng)")
            self.btn_email_retirement.clicked.connect(self.send_retirement_email)
            # chèn vào layout nút, đặt sau email báo cáo nhanh
            # chú ý: layout đã tạo sẵn, thêm vào cuối để tránh phá vỡ UI
            self.layout().itemAt(3).addWidget(self.btn_email_retirement) if hasattr(self, 'layout') else None
        self.btn_email_retirement.setEnabled(is_admin)
        is_mgr = role in ('admin','hr','unit_manager')
        self.btn_work.setEnabled(is_mgr)
        self.btn_contract_add.setEnabled(is_mgr)
        self.btn_contract_export.setEnabled(is_mgr)
        # Xuất lịch sử lương theo danh sách lọc: chỉ admin/hr và quản lý đơn vị (sẽ lọc theo đơn vị)
        self.btn_salary_export_filtered.setEnabled(is_mgr)
        # Gửi báo cáo nhanh (email): chỉ admin/hr
        self.btn_email_report.setEnabled(role in ('admin','hr'))

    def drain_notifications(self):
        # Lấy thông báo từ scheduler và hiển thị popup
        try:
            while not NOTIFY_QUEUE.empty():
                title, message = NOTIFY_QUEUE.get_nowait()
                QMessageBox.information(self, title, message)
        except Exception:
            pass

    def current_person(self):
        item = self.list.currentItem()
        if not item:
            return None, None
        code = item.text().split(" - ")[-1]
        db = SessionLocal()
        p = db.query(Person).filter_by(code=code).first()
        return db, p

    def export_selected(self):
        item = self.list.currentItem()
        if not item:
            QMessageBox.information(self, "Chưa chọn", "Vui lòng chọn một nhân sự trong danh sách")
            return
        db, person = self.current_person()
        try:
            if not person:
                QMessageBox.critical(self, "Lỗi", "Không tìm thấy nhân sự đã chọn")
                return
            Path("exports").mkdir(exist_ok=True)
            file_path = Path("exports") / f"trich_ngang_{person.code}.docx"
            # Nếu có template, render template; nếu không, fallback bảng mặc định
            tpl = Path("templates") / "trich_ngang_template.docx"
            context = {
                "full_name": person.full_name or "",
                "code": person.code or "",
                "dob": person.dob.isoformat() if person.dob else "",
                "gender": person.gender or "",
                "ethnicity": person.ethnicity or "",
                "religion": person.religion or "",
                "hometown": person.hometown or "",
                "position": person.position.name if person.position else "",
                "unit": person.unit.name if person.unit else "",
                "party_joined_date": person.party_joined_date.isoformat() if person.party_joined_date else "",
                "llct_level": person.llct_level or "",
                "professional_level": person.professional_level or "",
                "status": person.status or "",
                "phone": person.phone or "",
                "email": person.email or "",
            }
            if tpl.exists():
                render_docx_template(str(tpl), context, str(file_path))
            else:
                doc = Document()
                doc.add_heading("Trích ngang nhân sự", level=1)
                rows = [
                    ("Họ và tên", context["full_name"]),
                    ("Mã", context["code"]),
                    ("Ngày sinh", context["dob"]),
                    ("Giới tính", context["gender"]),
                    ("Dân tộc", context["ethnicity"]),
                    ("Tôn giáo", context["religion"]),
                    ("Quê quán", context["hometown"]),
                    ("Chức danh", context["position"]),
                    ("Đơn vị", context["unit"]),
                    ("Ngày vào Đảng", context["party_joined_date"]),
                    ("LLCT", context["llct_level"]),
                    ("Trình độ chuyên môn", context["professional_level"]),
                    ("Tình trạng công tác", context["status"]),
                    ("Điện thoại", context["phone"]),
                    ("Email", context["email"]),
                ]
                table = doc.add_table(rows=1, cols=2)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Trường"
                hdr_cells[1].text = "Giá trị"
                for k, v in rows:
                    row_cells = table.add_row().cells
                    row_cells[0].text = k
                    row_cells[1].text = v
                doc.save(str(file_path))
            # Audit
            try:
                log_action(db, self.current_user.get('id'), 'export_profile', 'Person', person.id, f"file={file_path}")
            except Exception:
                pass
            QMessageBox.information(self, "Thành công", f"Đã xuất: {file_path}")
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", str(e))
        finally:
            if db:
                db.close()

    def show_due(self):
        from datetime import date
        start, end = quarter_window(date.today())
        db = SessionLocal()
        try:
            items = list_due_in_window(db, start, end)
            if not items:
                QMessageBox.information(self, "Kết quả", "Không có nhân sự đến hạn trong quý này")
                return
            Path("exports").mkdir(exist_ok=True)
            xlsx = Path("exports") / f"nang_luong_quy_{end.year}_Q{((end.month-1)//3)+1}.xlsx"
            export_due_to_excel(items, str(xlsx), template_name=self.get_template_for('salary_due'), username=self.current_user.get('username'))
            # Audit
            try:
                log_action(db, self.current_user.get('id'), 'export_salary_due', 'Salary', None, f"file={xlsx}")
            except Exception:
                pass
            QMessageBox.information(self, "Thành công", f"Đã xuất danh sách: {xlsx}")
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", str(e))
        finally:
            db.close()

    def check_appointment(self):
        from .appointment import check_appointment_eligibility
        from PySide6.QtWidgets import QInputDialog
        db, person = self.current_person()
        try:
            if not person:
                QMessageBox.information(self, "Chưa chọn", "Chọn một nhân sự trước")
                return
            target, ok = QInputDialog.getText(self, "Vị trí bổ nhiệm", "Nhập vị trí cần kiểm tra:")
            if not ok or not target:
                return
            ok_eligible, reasons = check_appointment_eligibility(db, person, target)
            if ok_eligible:
                QMessageBox.information(self, "Kết quả", f"Đủ điều kiện bổ nhiệm vị trí: {target}")
            else:
                QMessageBox.warning(self, "Chưa đạt", "\n".join(reasons) or "Chưa đủ điều kiện")
            # Audit
            try:
                details = f"target={target}; ok={ok_eligible}; reasons={';'.join(reasons)}"
                log_action(db, self.current_user.get('id'), 'check_appointment', 'Appointment', None, details)
            except Exception:
                pass
        finally:
            if db:
                db.close()

    def add_work_process(self):
        from PySide6.QtWidgets import QInputDialog
        from datetime import date
        db, person = self.current_person()
        try:
            if not person:
                QMessageBox.information(self, "Chưa chọn", "Chọn một nhân sự trước")
                return
            unit, ok = QInputDialog.getText(self, "Đơn vị", "Nhập đơn vị:")
            if not ok or not unit:
                return
            pos, ok = QInputDialog.getText(self, "Chức vụ", "Nhập chức vụ:")
            if not ok or not pos:
                return
            from .models import WorkProcess
            # đơn giản: lấy start_date = ngày hiện tại
            wp = WorkProcess(person_id=person.id, unit=unit, position=pos, start_date=date.today())
            db.add(wp)
            db.commit()
            # Audit
            try:
                log_action(db, self.current_user.get('id'), 'add_work_process', 'WorkProcess', wp.id, f"unit={unit};pos={pos}")
            except Exception:
                pass
            QMessageBox.information(self, "Thành công", "Đã thêm quá trình công tác")
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", str(e))
        finally:
            if db:
                db.close()

    def open_detail(self):
        # Mở dialog chi tiết cho nhân sự đang chọn
        db, person = self.current_person()
        try:
            if not person:
                QMessageBox.information(self, "Chưa chọn", "Chọn một nhân sự trước")
                return
            dlg = PersonDetailDialog(person.id, self)
            dlg.exec()
        finally:
            if db:
                db.close()

    def export_work_process(self):
        from .models import WorkProcess
        db, person = self.current_person()
        try:
            if not person:
                QMessageBox.information(self, "Chưa chọn", "Chọn một nhân sự trước")
                return
            wps = db.query(WorkProcess).filter_by(person_id=person.id).order_by(WorkProcess.start_date).all()
            Path("exports").mkdir(exist_ok=True)
            file_path = Path("exports") / f"work_process_{person.code}.docx"
            doc = Document()
            doc.add_heading(f"Quá trình công tác - {person.full_name}", level=1)
            t = doc.add_table(rows=1, cols=3)
            t.rows[0].cells[0].text = "Từ ngày"
            t.rows[0].cells[1].text = "Đơn vị"
            t.rows[0].cells[2].text = "Chức vụ"
            for w in wps:
                row = t.add_row().cells
                row[0].text = (w.start_date.isoformat() + (" -> " + w.end_date.isoformat() if w.end_date else ""))
                row[1].text = w.unit or ""
                row[2].text = w.position or ""
            doc.save(str(file_path))
            # Audit
            try:
                log_action(db, self.current_user.get('id'), 'export_work_process', 'Person', person.id, f"file={file_path}")
            except Exception:
                pass
            QMessageBox.information(self, "Thành công", f"Đã xuất: {file_path}")
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", str(e))
        finally:
            if db:
                db.close()

    def export_salary_cover(self):
        # Xuất công văn rà soát nâng lương theo quý hiện tại
        from datetime import date
        from .docx_exports import export_salary_review_cover
        q = (date.today().month - 1)//3 + 1
        defaults = {
            'org_name': 'Cơ quan ABC',
            'quarter': str(q),
            'year': str(date.today().year),
            'date': date.today().isoformat(),
        }
        # Hỏi người dùng thông tin văn bản
        form_fields = [
            {'name': 'org_name', 'label': 'Đơn vị/Cơ quan', 'type': 'text', 'default': defaults['org_name']},
            {'name': 'quarter', 'label': 'Quý', 'type': 'select', 'default': defaults['quarter'], 'options': ['1','2','3','4']},
            {'name': 'year', 'label': 'Năm', 'type': 'text', 'default': defaults['year']},
            {'name': 'date', 'label': 'Ngày văn bản', 'type': 'date', 'default': defaults['date']},
        ]
        ctx = prompt_context(self, 'Thông tin công văn rà soát', form_fields)
        if not ctx:
            QMessageBox.information(self, "Đã hủy", "Đã hủy xuất công văn")
            return
        Path('exports').mkdir(exist_ok=True)
        out = Path('exports')/f"cong_van_ra_soat_quy_{ctx['year']}_Q{ctx['quarter']}.docx"
        try:
            export_salary_review_cover(ctx, str(out))
            # Optional PDF
            if self.btn_pdf_toggle.currentText() == 'PDF':
                from .templates import try_export_docx_to_pdf
                pdf_path = str(out).replace('.docx','.pdf')
                if try_export_docx_to_pdf(str(out), pdf_path):
                    out = Path(pdf_path)
            try:
                log_action(SessionLocal(), self.current_user.get('id'), 'export_salary_cover', 'Letter', None, f"file={out}")
            except Exception:
                pass
            QMessageBox.information(self, "Thành công", f"Đã xuất: {out}")
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", str(e))

    def export_retirement_letters(self):
        # Xuất văn bản nghỉ hưu cho nhân sự đã chọn (TB và QĐ)
        from datetime import date
        from .docx_exports import export_retirement_notification, export_retirement_decision
        db, person = self.current_person()
        try:
            if not person:
                QMessageBox.information(self, "Chưa chọn", "Chọn một nhân sự trước")
                return
            default_retire = (person.dob.replace(year=person.dob.year + (60 if (person.gender or 'Nam').lower().startswith('n') else 55)).isoformat() if person.dob else '')
            defaults = {
                'full_name': person.full_name or '',
                'code': person.code or '',
                'retirement_date': default_retire,
                'date': date.today().isoformat(),
            }
            fields = [
                {'name': 'full_name', 'label': 'Họ và tên', 'type': 'text', 'default': defaults['full_name']},
                {'name': 'code', 'label': 'Mã nhân sự', 'type': 'text', 'default': defaults['code']},
                {'name': 'retirement_date', 'label': 'Ngày nghỉ hưu', 'type': 'date', 'default': defaults['retirement_date']},
                {'name': 'date', 'label': 'Ngày văn bản', 'type': 'date', 'default': defaults['date']},
            ]
            ctx = prompt_context(self, 'Thông tin văn bản nghỉ hưu', fields)
            if not ctx:
                QMessageBox.information(self, "Đã hủy", "Đã hủy xuất văn bản")
                return
            Path('exports').mkdir(exist_ok=True)
            out1 = Path('exports')/f"thong_bao_nghi_huu_{person.code}.docx"
            out2 = Path('exports')/f"quyet_dinh_nghi_huu_{person.code}.docx"
            export_retirement_notification(ctx, str(out1))
            export_retirement_decision(ctx, str(out2))
            if self.btn_pdf_toggle.currentText() == 'PDF':
                from .templates import try_export_docx_to_pdf
                p1 = str(out1).replace('.docx','.pdf')
                p2 = str(out2).replace('.docx','.pdf')
                if try_export_docx_to_pdf(str(out1), p1):
                    out1 = Path(p1)
                if try_export_docx_to_pdf(str(out2), p2):
                    out2 = Path(p2)
            try:
                log_action(SessionLocal(), self.current_user.get('id'), 'export_retirement_letters', 'Letter', person.id, f"files={out1},{out2}")
            except Exception:
                pass
            QMessageBox.information(self, "Thành công", f"Đã xuất: {out1}\n{out2}")
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", str(e))
        finally:
            if db:
                db.close()

    def export_salary_letters(self):
        # Xuất Thông báo & Quyết định nâng lương cho nhân sự đã chọn nếu có đề xuất
        from datetime import date
        from .salary import compute_next_for_person
        from .docx_exports import export_salary_notification, export_salary_decision
        db, person = self.current_person()
        try:
            if not person:
                QMessageBox.information(self, "Chưa chọn", "Chọn một nhân sự trước")
                return
            info = compute_next_for_person(db, person, date.today())
            if not info:
                QMessageBox.information(self, "Không có đề xuất", "Nhân sự chưa đến hạn nâng lương")
                return
            effective = info.get('due_date')
            step = info.get('next_step') if info.get('type') == 'step' else ''
            coef = info.get('next_coef') if info.get('type') == 'step' else f"{info.get('allowance_percent','')}%"
            defaults = {
                'full_name': person.full_name or '',
                'code': person.code or '',
                'unit': person.unit.name if person.unit else '',
                'effective_date': effective.isoformat() if effective else date.today().isoformat(),
                'step': str(step) if step != '' else '',
                'coefficient': str(coef) if coef is not None else '',
            }
            fields = [
                {'name': 'full_name', 'label': 'Họ và tên', 'type': 'text', 'default': defaults['full_name']},
                {'name': 'code', 'label': 'Mã nhân sự', 'type': 'text', 'default': defaults['code']},
                {'name': 'unit', 'label': 'Đơn vị', 'type': 'text', 'default': defaults['unit']},
                {'name': 'effective_date', 'label': 'Ngày hiệu lực', 'type': 'date', 'default': defaults['effective_date']},
                {'name': 'step', 'label': 'Bậc', 'type': 'text', 'default': defaults['step']},
                {'name': 'coefficient', 'label': 'Hệ số/% vượt khung', 'type': 'text', 'default': defaults['coefficient']},
            ]
            ctx = prompt_context(self, 'Thông tin văn bản nâng lương', fields)
            if not ctx:
                QMessageBox.information(self, "Đã hủy", "Đã hủy xuất văn bản")
                return
            Path('exports').mkdir(exist_ok=True)
            out1 = Path('exports')/f"thong_bao_nang_luong_{person.code}.docx"
            out2 = Path('exports')/f"quyet_dinh_nang_luong_{person.code}.docx"
            export_salary_notification(ctx, str(out1))
            export_salary_decision(ctx, str(out2))
            if self.btn_pdf_toggle.currentText() == 'PDF':
                from .templates import try_export_docx_to_pdf
                p1 = str(out1).replace('.docx','.pdf')
                p2 = str(out2).replace('.docx','.pdf')
                if try_export_docx_to_pdf(str(out1), p1):
                    out1 = Path(p1)
                if try_export_docx_to_pdf(str(out2), p2):
                    out2 = Path(p2)
            try:
                log_action(SessionLocal(), self.current_user.get('id'), 'export_salary_letters', 'Letter', person.id, f"files={out1},{out2}")
            except Exception:
                pass
            QMessageBox.information(self, "Thành công", f"Đã xuất: {out1}\n{out2}")
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", str(e))
        finally:
            if db:
                db.close()

    def send_quarter_salary_email(self):
        # Ghi EmailLog tổng hợp cho nâng lương quý sau khi gửi (đã có ghi audit)
        role = (self.current_user.get('role') or '').lower()
        if role not in ('admin','hr'):
            QMessageBox.warning(self, "Không có quyền", "Chỉ admin/HR mới gửi email")
            return
        from datetime import date
        from pathlib import Path
        from .salary import list_due_in_window, export_due_to_excel
        from .mailer import send_email_with_attachment
        start, end = quarter_window(date.today())
        db = SessionLocal()
        try:
            items = list_due_in_window(db, start, end)
            if not items:
                QMessageBox.information(self, "Kết quả", "Không có nhân sự đến hạn trong quý này")
                return
            Path("exports").mkdir(exist_ok=True)
            q = ((end.month - 1)//3) + 1
            out = Path("exports") / f"nang_luong_quy_{end.year}_Q{q}.xlsx"
            export_due_to_excel(items, str(out), template_name=self.get_template_for('salary_due'), username=self.current_user.get('username'))
            subject = f"[HRMS] Danh sách đến hạn nâng lương - Q{q}/{end.year}"
            body = f"Có {len(items)} nhân sự đến hạn trong quý này. Tệp đính kèm: {out.name}"
            ok = send_email_with_attachment(subject, body, [str(out)])
            # Audit
            try:
                log_action(SessionLocal(), self.current_user.get('id'), 'email_salary_due', 'Salary', None, f"file={out};sent={ok};count={len(items)}")
                # EmailLog chi tiết
                from .db import SessionLocal as _SL
                from .models import EmailLog
                s = _SL()
                s.add(EmailLog(type='salary_due', unit_name=None, recipients='', subject=subject, body=body[:1000], attachments=str(out), status='sent' if ok else 'failed', user_id=self.current_user.get('id')))
                s.commit(); s.close()
            except Exception:
                pass
            if ok:
                QMessageBox.information(self, "Thành công", f"Đã gửi email kèm tệp: {out}")
            else:
                QMessageBox.warning(self, "Chưa gửi", "Không gửi được email. Kiểm tra cấu hình SMTP/ALERT_EMAILS")
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", str(e))
        finally:
            db.close()

    def send_quarter_salary_email_by_unit(self):
        # Gửi nâng lương theo từng đơn vị + ZIP tổng hợp nếu bật
        role = (self.current_user.get('role') or '').lower()
        if role not in ('admin','hr'):
            QMessageBox.warning(self, "Không có quyền", "Chỉ admin/HR mới gửi")
            return
        from datetime import date
        from pathlib import Path
        from .salary import list_due_in_window, export_due_to_excel
        from .mailer import send_email_with_attachment, get_recipients_for_unit, create_zip
        from .settings_service import get_setting
        start, end = quarter_window(date.today())
        db = SessionLocal()
        try:
            # Thu thập đơn vị có dữ liệu và người nhận
            from .models import Unit
            units = db.query(Unit).order_by(Unit.name).all()
            data = []  # (checkbox, unit_name, unit_id, count, recips)
            from PySide6.QtWidgets import QDialog, QVBoxLayout, QCheckBox, QPushButton, QLabel, QHBoxLayout
            dlg = QDialog(self); dlg.setWindowTitle("Gửi nâng lương theo đơn vị")
            v = QVBoxLayout(dlg)
            v.addWidget(QLabel(f"Quý hiện tại: Q{((end.month-1)//3)+1}/{end.year}"))
            checks = []
            for u in units:
                arr = list_due_in_window(db, start, end, unit_id=u.id)
                recips = get_recipients_for_unit(u.name)
                text = f"{u.name} - {len(arr)} người; {len(recips)} email"
                cb = QCheckBox(text)
                cb.setChecked(bool(arr and recips))
                v.addWidget(cb)
                checks.append(cb)
                data.append((cb, u.name, u.id, len(arr), recips))
            row = QHBoxLayout(); btn_ok = QPushButton("Gửi"); btn_cancel = QPushButton("Hủy"); btn_all = QPushButton("Chọn tất cả"); btn_none = QPushButton("Bỏ chọn")
            def sel_all():
                for c in checks: c.setChecked(True)
            def sel_none():
                for c in checks: c.setChecked(False)
            btn_all.clicked.connect(sel_all); btn_none.clicked.connect(sel_none)
            row.addWidget(btn_all); row.addWidget(btn_none); row.addWidget(btn_ok); row.addWidget(btn_cancel)
            v.addLayout(row)
            zip_cb = QCheckBox("Gửi kèm ZIP tổng hợp");
            try:
                zip_cb.setChecked(((get_setting('SEND_SUMMARY_ZIP','0') or '0').strip().lower() in ('1','true','yes')))
            except Exception:
                zip_cb.setChecked(False)
            v.addWidget(zip_cb)

            def do_send():
                try:
                    Path('exports').mkdir(exist_ok=True)
                    q = ((end.month - 1)//3) + 1
                    sent_files = []
                    sent = 0; failed = 0
                    for cb, unit_name, unit_id, cnt, recips in data:
                        if not cb.isChecked() or cnt <= 0 or not recips:
                            continue
                        arr = list_due_in_window(db, start, end, unit_id=unit_id)
                        if not arr:
                            continue
                        out = Path('exports')/f"nang_luong_quy_{end.year}_Q{q}_{unit_name.replace(' ', '_')}.xlsx"
                        export_due_to_excel(arr, str(out), template_name=self.get_template_for('salary_due'), username=self.current_user.get('username'))
                        ok = send_email_with_attachment(f"[HRMS] Nâng lương Q{q}/{end.year} - {unit_name}", f"Đính kèm danh sách cho {unit_name}", [str(out)], to=recips)
                        sent_files.append(str(out))
                        try:
                            # Audit
                            log_action(SessionLocal(), self.current_user.get('id'), 'email_salary_due_unit', 'Salary', None, f"unit={unit_name};file={out};sent={ok};count={len(arr)}")
                            # EmailLog chi tiết (mask recipients)
                            masked = []
                            for r in recips:
                                r = str(r or '').strip()
                                if '@' in r:
                                    masked.append('***@' + r.split('@',1)[1])
                                elif r:
                                    masked.append('***')
                            from .models import EmailLog
                            s = SessionLocal();
                            s.add(EmailLog(type='salary_due', unit_name=unit_name, recipients=", ".join(masked)[:1000], subject=f"Nâng lương Q{q}/{end.year} - {unit_name}", body=f"Gửi danh sách cho {unit_name}", attachments=str(out), status='sent' if ok else 'failed', user_id=self.current_user.get('id')))
                            s.commit(); s.close()
                        except Exception:
                            pass
                        if ok: sent += 1
                        else: failed += 1
                    # ZIP tổng hợp nếu checkbox bật
                    try:
                        if zip_cb.isChecked() and sent_files:
                            zip_path = Path('exports')/f"nang_luong_quy_{end.year}_Q{q}_by_unit.zip"
                            if create_zip(sent_files, str(zip_path)):
                                send_email_with_attachment(f"[HRMS] Nâng lương Q{q}/{end.year} (ZIP tổng hợp)", f"ZIP tổng hợp danh sách nâng lương theo đơn vị Q{q}/{end.year}", [str(zip_path)])
                    except Exception:
                        pass
                    QMessageBox.information(dlg, "Kết quả", f"Gửi thành công: {sent}\nThất bại: {failed}")
                    dlg.accept()
                except Exception as ex:
                    QMessageBox.critical(dlg, "Lỗi", str(ex))
            btn_ok.clicked.connect(do_send)
            btn_cancel.clicked.connect(dlg.reject)
            dlg.resize(520, 600)
            dlg.exec()
        finally:
            db.close()

    def send_insurance_email_by_unit(self):
        # Gửi BHXH tháng trước theo từng đơn vị + ZIP tổng hợp
        role = (self.current_user.get('role') or '').lower()
        if role not in ('admin','hr'):
            QMessageBox.warning(self, "Không có quyền", "Chỉ admin/HR mới gửi")
            return
        from datetime import date
        from calendar import monthrange
        from pathlib import Path
        from .insurance import export_insurance_to_excel
        from .mailer import send_email_with_attachment, get_recipients_for_unit, create_zip
        from .settings_service import get_setting
        today = date.today()
        prev_month = (today.month - 2) % 12 + 1
        prev_year = today.year - 1 if today.month == 1 else today.year
        start = date(prev_year, prev_month, 1)
        end = date(prev_year, prev_month, monthrange(prev_year, prev_month)[1])
        db = SessionLocal()
        try:
            from .models import Unit
            units = db.query(Unit).order_by(Unit.name).all()
            from PySide6.QtWidgets import QDialog, QVBoxLayout, QCheckBox, QPushButton, QLabel, QHBoxLayout
            dlg = QDialog(self); dlg.setWindowTitle("Gửi BHXH theo đơn vị")
            v = QVBoxLayout(dlg)
            v.addWidget(QLabel(f"Tháng: {prev_month:02d}/{prev_year}"))
            data = []
            checks = []
            for u in units:
                recips = get_recipients_for_unit(u.name)
                text = f"{u.name} - {len(recips)} email"
                cb = QCheckBox(text); cb.setChecked(bool(recips))
                v.addWidget(cb)
                checks.append(cb)
                data.append((cb, u.name, u.id, recips))
            row = QHBoxLayout(); btn_ok = QPushButton("Gửi"); btn_cancel = QPushButton("Hủy"); btn_all = QPushButton("Chọn tất cả"); btn_none = QPushButton("Bỏ chọn")
            def sel_all():
                for c in checks: c.setChecked(True)
            def sel_none():
                for c in checks: c.setChecked(False)
            btn_all.clicked.connect(sel_all); btn_none.clicked.connect(sel_none)
            row.addWidget(btn_all); row.addWidget(btn_none); row.addWidget(btn_ok); row.addWidget(btn_cancel)
            v.addLayout(row)
            zip_cb = QCheckBox("Gửi kèm ZIP tổng hợp");
            try:
                zip_cb.setChecked(((get_setting('SEND_SUMMARY_ZIP','0') or '0').strip().lower() in ('1','true','yes')))
            except Exception:
                zip_cb.setChecked(False)
            v.addWidget(zip_cb)
            def do_send():
                try:
                    Path('exports').mkdir(exist_ok=True)
                    sent_files = []
                    sent = 0; failed = 0
                    for cb, unit_name, unit_id, recips in data:
                        if not cb.isChecked() or not recips:
                            continue
                        out = Path('exports')/f"bhxh_{prev_year}_{prev_month:02d}_{unit_name.replace(' ', '_')}.xlsx"
                        export_insurance_to_excel(db, start, end, str(out), template_name=self.get_template_for('bhxh'), username=self.current_user.get('username'), unit_id=unit_id)
                        ok = send_email_with_attachment(f"[HRMS] BHXH {prev_month:02d}/{prev_year} - {unit_name}", f"Đính kèm BHXH {prev_month:02d}/{prev_year} - {unit_name}", [str(out)], to=recips)
                        sent_files.append(str(out))
                        try:
                            # Audit + EmailLog
                            log_action(SessionLocal(), self.current_user.get('id'), 'email_insurance_unit', 'InsuranceEvent', None, f"unit={unit_name};file={out};sent={ok}")
                            masked = []
                            for r in recips:
                                r = str(r or '').strip()
                                if '@' in r:
                                    masked.append('***@' + r.split('@',1)[1])
                                elif r:
                                    masked.append('***')
                            from .models import EmailLog
                            s = SessionLocal(); s.add(EmailLog(type='bhxh_monthly', unit_name=unit_name, recipients=", ".join(masked)[:1000], subject=f"BHXH {prev_month:02d}/{prev_year} - {unit_name}", body=f"BHXH {prev_month:02d}/{prev_year}", attachments=str(out), status='sent' if ok else 'failed', user_id=self.current_user.get('id'))); s.commit(); s.close()
                        except Exception: pass
                        if ok: sent += 1
                        else: failed += 1
                    # ZIP tổng hợp nếu checkbox bật
                    try:
                        if zip_cb.isChecked() and sent_files:
                            zip_path = Path('exports')/f"bhxh_{prev_year}_{prev_month:02d}_by_unit.zip"
                            if create_zip(sent_files, str(zip_path)):
                                send_email_with_attachment(f"[HRMS] BHXH {prev_month:02d}/{prev_year} (ZIP tổng hợp)", f"ZIP tổng hợp BHXH {prev_month:02d}/{prev_year}", [str(zip_path)])
                    except Exception:
                        pass
                    QMessageBox.information(dlg, "Kết quả", f"Gửi thành công: {sent}\nThất bại: {failed}")
                    dlg.accept()
                except Exception as ex:
                    QMessageBox.critical(dlg, "Lỗi", str(ex))
            btn_ok.clicked.connect(do_send)
            btn_cancel.clicked.connect(dlg.reject)
            dlg.resize(520, 600)
            dlg.exec()
        finally:
            db.close()

    def send_contracts_expiring_email_by_unit(self):
        # Gửi HĐ sắp hết hạn theo đơn vị + ZIP tổng hợp
        role = (self.current_user.get('role') or '').lower()
        if role not in ('admin','hr'):
            QMessageBox.warning(self, "Không có quyền", "Chỉ admin/HR mới gửi")
            return
        from datetime import date, timedelta
        from pathlib import Path
        from .contracts import export_contracts_expiring_to_excel
        from .mailer import send_email_with_attachment, get_recipients_for_unit, create_zip
        from .settings_service import get_setting
        today = date.today()
        days = int(get_setting('CONTRACT_ALERT_DAYS','30') or '30')
        end = today + timedelta(days=days)
        db = SessionLocal()
        try:
            from .models import Unit
            units = db.query(Unit).order_by(Unit.name).all()
            from PySide6.QtWidgets import QDialog, QVBoxLayout, QCheckBox, QPushButton, QLabel, QHBoxLayout, QLineEdit
            dlg = QDialog(self); dlg.setWindowTitle("Gửi HĐ sắp hết hạn theo đơn vị")
            v = QVBoxLayout(dlg)
            v.addWidget(QLabel(f"Trong {days} ngày tới"))
            data = []
            checks = []
            for u in units:
                recips = get_recipients_for_unit(u.name)
                text = f"{u.name} - {len(recips)} email"
                cb = QCheckBox(text); cb.setChecked(bool(recips))
                v.addWidget(cb)
                checks.append(cb)
                data.append((cb, u.name, u.id, recips))
            row = QHBoxLayout(); btn_ok = QPushButton("Gửi"); btn_cancel = QPushButton("Hủy"); btn_all = QPushButton("Chọn tất cả"); btn_none = QPushButton("Bỏ chọn")
            def sel_all():
                for c in checks: c.setChecked(True)
            def sel_none():
                for c in checks: c.setChecked(False)
            btn_all.clicked.connect(sel_all); btn_none.clicked.connect(sel_none)
            row.addWidget(btn_all); row.addWidget(btn_none); row.addWidget(btn_ok); row.addWidget(btn_cancel)
            v.addLayout(row)
            zip_cb = QCheckBox("Gửi kèm ZIP tổng hợp");
            try:
                zip_cb.setChecked(((get_setting('SEND_SUMMARY_ZIP','0') or '0').strip().lower() in ('1','true','yes')))
            except Exception:
                zip_cb.setChecked(False)
            v.addWidget(zip_cb)
            def do_send():
                try:
                    Path('exports').mkdir(exist_ok=True)
                    sent_files = []
                    sent = 0; failed = 0
                    for cb, unit_name, unit_id, recips in data:
                        if not cb.isChecked() or not recips:
                            continue
                        out = Path('exports')/f"contracts_expiring_{today.isoformat()}_{days}d_{unit_name.replace(' ', '_')}.xlsx"
                        export_contracts_expiring_to_excel(db, today, end, str(out), template_name=self.get_template_for('contracts'), username=self.current_user.get('username'), unit_id=unit_id)
                        ok = send_email_with_attachment(f"[HRMS] HĐ sắp hết hạn ({days} ngày) - {unit_name}", f"Đính kèm HĐ sắp hết hạn - {unit_name}", [str(out)], to=recips)
                        sent_files.append(str(out))
                        try:
                            # Audit + EmailLog
                            log_action(SessionLocal(), self.current_user.get('id'), 'email_contracts_expiring_unit', 'Contract', None, f"unit={unit_name};file={out};sent={ok}")
                            masked = []
                            for r in recips:
                                r = str(r or '').strip()
                                if '@' in r:
                                    masked.append('***@' + r.split('@',1)[1])
                                elif r:
                                    masked.append('***')
                            from .models import EmailLog
                            s = SessionLocal(); s.add(EmailLog(type='contracts_expiring', unit_name=unit_name, recipients=", ".join(masked)[:1000], subject=f"HĐ sắp hết hạn ({days}d) - {unit_name}", body=f"HĐ sắp hết hạn - {unit_name}", attachments=str(out), status='sent' if ok else 'failed', user_id=self.current_user.get('id'))); s.commit(); s.close()
                        except Exception: pass
                        if ok: sent += 1
                        else: failed += 1
                    # ZIP tổng hợp nếu checkbox bật
                    try:
                        if zip_cb.isChecked() and sent_files:
                            zip_path = Path('exports')/f"contracts_expiring_{today.isoformat()}_{days}d_by_unit.zip"
                            if create_zip(sent_files, str(zip_path)):
                                send_email_with_attachment(f"[HRMS] HĐ sắp hết hạn (ZIP tổng hợp)", 'ZIP tổng hợp HĐ sắp hết hạn các đơn vị', [str(zip_path)])
                    except Exception:
                        pass
                    QMessageBox.information(dlg, "Kết quả", f"Gửi thành công: {sent}\nThất bại: {failed}")
                    dlg.accept()
                except Exception as ex:
                    QMessageBox.critical(dlg, "Lỗi", str(ex))
            btn_ok.clicked.connect(do_send)
            btn_cancel.clicked.connect(dlg.reject)
            dlg.resize(520, 600)
            dlg.exec()
        finally:
            db.close()

    def send_retirement_email(self):
        # Gửi email danh sách nghỉ hưu (6/3 tháng) kèm Excel 2 sheet
        role = (self.current_user.get('role') or '').lower()
        if role not in ('admin','hr'):
            QMessageBox.warning(self, "Không có quyền", "Chỉ admin/HR mới gửi")
            return
        from datetime import date
        from pathlib import Path
        from .reporting import export_retirement_alerts_to_excel
        from .retirement import calculate_retirement_date
        from .mailer import send_email_with_attachment
        db = SessionLocal()
        try:
            today = date.today()
            six = date(today.year + (today.month + 6 - 1) // 12, ((today.month + 6 - 1) % 12) + 1, today.day)
            three = date(today.year + (today.month + 3 - 1) // 12, ((today.month + 3 - 1) % 12) + 1, today.day)
            persons = db.query(Person).all()
            list6 = [p for p in persons if calculate_retirement_date(p) == six]
            list3 = [p for p in persons if calculate_retirement_date(p) == three]
            if not list6 and not list3:
                QMessageBox.information(self, "Kết quả", "Chưa có người trùng mốc 6/3 tháng kể từ hôm nay")
                return
            Path('exports').mkdir(exist_ok=True)
            out = Path('exports')/f"nghi_huu_thong_bao_{today.isoformat()}.xlsx"
            export_retirement_alerts_to_excel(db, list6, list3, str(out))
            subject = f"[HRMS] Danh sách nghỉ hưu (6/3 tháng) {today.isoformat()}"
            body = "Đính kèm danh sách nghỉ hưu dự kiến theo mốc 6 và 3 tháng."
            ok = send_email_with_attachment(subject, body, [str(out)])
            try:
                log_action(SessionLocal(), self.current_user.get('id'), 'email_retirement_alerts', 'Retirement', None, f"file={out};sent={ok};six={len(list6)};three={len(list3)}")
                # EmailLog chi tiết
                from .db import SessionLocal as _SL
                from .models import EmailLog
                s = _SL(); s.add(EmailLog(type='retirement', unit_name=None, recipients='', subject=subject, body=body[:1000], attachments=str(out), status='sent' if ok else 'failed', user_id=self.current_user.get('id'))); s.commit(); s.close()
            except Exception:
                pass
            if ok:
                QMessageBox.information(self, "Thành công", f"Đã gửi email: {out}")
            else:
                QMessageBox.warning(self, "Chưa gửi", "Không gửi được email. Kiểm tra cấu hình SMTP/ALERT_EMAILS")
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", str(e))
        finally:
            db.close()

    def quick_report(self):
        from datetime import date
        from .reporting import compute_annual_summary, compute_demographics, export_report_to_excel
        year = date.today().year
        db = SessionLocal()
        try:
            summary = compute_annual_summary(db, year)
            demo = compute_demographics(db, date.today())
            Path("exports").mkdir(exist_ok=True)
            file_path = Path("exports") / f"bao_cao_nhanh_{year}.xlsx"
            export_report_to_excel(summary, demo, str(file_path))
            # Audit
            try:
                log_action(db, self.current_user.get('id'), 'export_quick_report', 'Report', None, f"file={file_path}")
            except Exception:
                pass
            QMessageBox.information(self, "Thành công", f"Đã xuất: {file_path}")
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", str(e))
        finally:
            db.close()

    def send_quick_report_email(self):
        # Gửi báo cáo nhanh qua email (đính kèm file Excel)
        role = (self.current_user.get('role') or '').lower()
        if role not in ('admin','hr'):
            QMessageBox.warning(self, "Không có quyền", "Chỉ admin/HR mới gửi")
            return
        from datetime import date
        from pathlib import Path
        from .reporting import compute_annual_summary, compute_demographics, export_report_to_excel
        from .mailer import send_email_with_attachment
        year = date.today().year
        db = SessionLocal()
        try:
            summary = compute_annual_summary(db, year)
            demo = compute_demographics(db, date.today())
            Path("exports").mkdir(exist_ok=True)
            file_path = Path("exports") / f"bao_cao_nhanh_{year}.xlsx"
            export_report_to_excel(summary, demo, str(file_path))
            subject = f"[HRMS] Báo cáo nhanh {year}"
            body = "Đính kèm báo cáo nhanh tổng hợp và cơ cấu nhân sự."
            ok = send_email_with_attachment(subject, body, [str(file_path)])
            try:
                log_action(SessionLocal(), self.current_user.get('id'), 'email_quick_report', 'Report', None, f"file={file_path};sent={ok}")
                # EmailLog chi tiết
                from .db import SessionLocal as _SL
                from .models import EmailLog
                s = _SL(); s.add(EmailLog(type='quick_report', unit_name=None, recipients='', subject=subject, body=body[:1000], attachments=str(file_path), status='sent' if ok else 'failed', user_id=self.current_user.get('id'))); s.commit(); s.close()
            except Exception:
                pass
            if ok:
                QMessageBox.information(self, "Thành công", f"Đã gửi email: {file_path}")
            else:
                QMessageBox.warning(self, "Chưa gửi", "Không gửi được email. Kiểm tra cấu hình SMTP/ALERT_EMAILS")
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", str(e))
        finally:
            db.close()

    def export_salary_histories_filtered(self):
        # Xuất lịch sử lương cho danh sách đang hiển thị (lọc)
        from datetime import datetime
        from .salary import export_salary_histories_for_people
        db = SessionLocal()
        try:
            role = (self.current_user.get('role') or '').lower()
            # Lấy danh sách code từ list widget (đang hiển thị theo bộ lọc)
            codes = []
            for i in range(self.list.count()):
                item = self.list.item(i)
                if not item:
                    continue
                code = item.text().split(" - ")[-1]
                codes.append(code)
            if not codes:
                QMessageBox.information(self, "Không có dữ liệu", "Không có nhân sự trong danh sách để xuất")
                return
            from .models import Person
            q = db.query(Person).filter(Person.code.in_(codes))
            # RBAC: nếu quản lý đơn vị, chỉ cho phép xuất người cùng đơn vị
            if role == 'unit_manager':
                only_unit_id = self.current_user.get('unit_id')
                if only_unit_id:
                    q = q.filter(Person.unit_id == only_unit_id)
            people = q.order_by(Person.full_name).all()
            if not people:
                QMessageBox.information(self, "Không có dữ liệu", "Không có nhân sự phù hợp quyền để xuất")
                return
            Path("exports").mkdir(exist_ok=True)
            ts = datetime.now().strftime('%Y%m%d_%H%M%S')
            out = Path("exports")/f"salary_histories_{ts}.xlsx"
            export_salary_histories_for_people(db, people, str(out), template_name=self.get_template_for('salary_histories'), username=self.current_user.get('username'))
            try:
                log_action(db, self.current_user.get('id'), 'export_salary_histories', 'Person', None, f"count={len(people)};file={out}")
            except Exception:
                pass
            QMessageBox.information(self, "Thành công", f"Đã xuất: {out}")
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", str(e))
        finally:
            db.close()

    def open_settings(self):
        # Trang cấu hình tối giản: SMTP và alert emails
        role = (self.current_user.get('role') or '').lower()
        if role not in ('admin','hr'):
            QMessageBox.warning(self, "Không có quyền", "Chỉ admin/HR mới truy cập cấu hình")
            return
        from PySide6.QtWidgets import QDialog, QFormLayout
        from .settings_service import get_setting, set_setting
        dlg = QDialog(self); dlg.setWindowTitle("Cấu hình hệ thống")
        f = QFormLayout(dlg)
        smtp_host = QLineEdit(get_setting('SMTP_HOST','') or '')
        smtp_port = QLineEdit(get_setting('SMTP_PORT','587') or '587')
        smtp_user = QLineEdit(get_setting('SMTP_USER','') or '')
        smtp_pass = QLineEdit(get_setting('SMTP_PASSWORD','') or ''); smtp_pass.setEchoMode(QLineEdit.Password)
        alert_emails = QLineEdit(get_setting('ALERT_EMAILS','') or '')
        org_name = QLineEdit(get_setting('ORG_NAME','') or '')
        date_fmt = QLineEdit(get_setting('XLSX_DATE_FORMAT','DD/MM/YYYY') or 'DD/MM/YYYY')
        coef_fmt = QLineEdit(get_setting('XLSX_NUMBER_FORMAT_COEF','0.00') or '0.00')
        freeze_global = QLineEdit(get_setting('XLSX_FREEZE_COL:GLOBAL','A') or 'A')
        freeze_bhxh = QLineEdit(get_setting('XLSX_FREEZE_COL:bhxh','') or '')
        freeze_contracts = QLineEdit(get_setting('XLSX_FREEZE_COL:contracts','') or '')
        freeze_salary_due = QLineEdit(get_setting('XLSX_FREEZE_COL:salary_due','') or '')
        freeze_salary_history = QLineEdit(get_setting('XLSX_FREEZE_COL:salary_history','') or '')
        freeze_salary_histories = QLineEdit(get_setting('XLSX_FREEZE_COL:salary_histories','') or '')
        f.addRow("SMTP_HOST", smtp_host)
        f.addRow("SMTP_PORT", smtp_port)
        f.addRow("SMTP_USER", smtp_user)
        f.addRow("SMTP_PASSWORD", smtp_pass)
        f.addRow("ALERT_EMAILS", alert_emails)
        f.addRow("ORG_NAME", org_name)
        # Các cờ email/ZIP tổng hợp
        subject_prefix = QLineEdit(get_setting('EMAIL_SUBJECT_PREFIX','') or '')
        unit_emails = QLineEdit(get_setting('UNIT_EMAILS','') or '')
        summary_zip = QLineEdit(get_setting('SEND_SUMMARY_ZIP','0') or '0')
        contract_alert_days = QLineEdit(get_setting('CONTRACT_ALERT_DAYS','30') or '30')
        export_ttl_days = QLineEdit(get_setting('EXPORT_TTL_DAYS','30') or '30')
        email_log_ttl = QLineEdit(get_setting('EMAIL_LOG_TTL_DAYS','365') or '365')
        f.addRow("EMAIL_SUBJECT_PREFIX", subject_prefix)
        f.addRow("UNIT_EMAILS", unit_emails)
        f.addRow("SEND_SUMMARY_ZIP (1/0)", summary_zip)
        f.addRow("CONTRACT_ALERT_DAYS", contract_alert_days)
        f.addRow("RETRY_COUNT", QLineEdit(get_setting('RETRY_COUNT','2') or '2'))
        f.addRow("RETRY_DELAY", QLineEdit(get_setting('RETRY_DELAY','10') or '10'))
        f.addRow("EXPORT_TTL_DAYS", export_ttl_days)
        f.addRow("EMAIL_LOG_TTL_DAYS", email_log_ttl)
        f.addRow("XLSX_DATE_FORMAT", date_fmt)
        f.addRow("XLSX_NUMBER_FORMAT_COEF", coef_fmt)
        f.addRow("XLSX_FREEZE_COL:GLOBAL", freeze_global)
        f.addRow("XLSX_FREEZE_COL:bhxh", freeze_bhxh)
        f.addRow("XLSX_FREEZE_COL:contracts", freeze_contracts)
        f.addRow("XLSX_FREEZE_COL:salary_due", freeze_salary_due)
        f.addRow("XLSX_FREEZE_COL:salary_history", freeze_salary_history)
        f.addRow("XLSX_FREEZE_COL:salary_histories", freeze_salary_histories)
        btn_ok = QPushButton("Lưu"); btn_cancel = QPushButton("Đóng"); btn_unit_emails = QPushButton("Email theo đơn vị…"); btn_unit_emails_db = QPushButton("QL email theo đơn vị (DB)…"); btn_test_smtp = QPushButton("Test SMTP")
        row = QHBoxLayout(); row.addWidget(btn_ok); row.addWidget(btn_cancel); row.addWidget(btn_unit_emails); row.addWidget(btn_unit_emails_db); row.addWidget(btn_test_smtp)
        f.addRow(row)

        def save():
            set_setting('SMTP_HOST', smtp_host.text())
            set_setting('SMTP_PORT', smtp_port.text())
            set_setting('SMTP_USER', smtp_user.text())
            set_setting('SMTP_PASSWORD', smtp_pass.text())
            set_setting('ALERT_EMAILS', alert_emails.text())
            set_setting('ORG_NAME', org_name.text())
            set_setting('EMAIL_SUBJECT_PREFIX', subject_prefix.text())
            set_setting('UNIT_EMAILS', unit_emails.text())
            set_setting('SEND_SUMMARY_ZIP', (summary_zip.text() or '0'))
            set_setting('CONTRACT_ALERT_DAYS', (contract_alert_days.text() or '30'))
            # RETRY_COUNT/DELAY: đọc từ form theo label
            try:
                from PySide6.QtWidgets import QLineEdit
                # QFormLayout order known: locate by label text
                for i in range(f.rowCount()):
                    label_item = f.itemAt(i, QFormLayout.LabelRole)
                    field_item = f.itemAt(i, QFormLayout.FieldRole)
                    if not label_item or not field_item: continue
                    wlabel = label_item.widget()
                    wfield = field_item.widget()
                    if not wlabel or not wfield: continue
                    name = getattr(wlabel, 'text', lambda: '')()
                    if name == 'RETRY_COUNT' and isinstance(wfield, QLineEdit):
                        set_setting('RETRY_COUNT', (wfield.text() or '2'))
                    if name == 'RETRY_DELAY' and isinstance(wfield, QLineEdit):
                        set_setting('RETRY_DELAY', (wfield.text() or '10'))
            except Exception:
                pass
            set_setting('EXPORT_TTL_DAYS', (export_ttl_days.text() or '30'))
            set_setting('EMAIL_LOG_TTL_DAYS', (email_log_ttl.text() or '365'))
            set_setting('XLSX_DATE_FORMAT', (date_fmt.text() or 'DD/MM/YYYY'))
            set_setting('XLSX_NUMBER_FORMAT_COEF', (coef_fmt.text() or '0.00'))
            set_setting('XLSX_FREEZE_COL:GLOBAL', (freeze_global.text() or 'A').upper())
            if freeze_bhxh.text(): set_setting('XLSX_FREEZE_COL:bhxh', freeze_bhxh.text().upper())
            if freeze_contracts.text(): set_setting('XLSX_FREEZE_COL:contracts', freeze_contracts.text().upper())
            if freeze_salary_due.text(): set_setting('XLSX_FREEZE_COL:salary_due', freeze_salary_due.text().upper())
            if freeze_salary_history.text(): set_setting('XLSX_FREEZE_COL:salary_history', freeze_salary_history.text().upper())
            if freeze_salary_histories.text(): set_setting('XLSX_FREEZE_COL:salary_histories', freeze_salary_histories.text().upper())
            QMessageBox.information(dlg, "Đã lưu", "Lưu cấu hình thành công")

        def edit_unit_emails():
            # Dialog cấu hình email theo đơn vị
            udlg = QDialog(dlg); udlg.setWindowTitle("Email theo đơn vị");
            from PySide6.QtWidgets import QFormLayout
            uf = QFormLayout(udlg)
            # Fetch units
            from .models import Unit
            from .db import SessionLocal as _SL
            dbu = _SL()
            try:
                units = dbu.query(Unit).order_by(Unit.name).all()
            finally:
                dbu.close()
            # Parse existing mapping
            import json
            raw = get_setting('UNIT_EMAILS','') or ''
            mapping: dict[str, list[str]] = {}
            if raw.strip():
                try:
                    obj = json.loads(raw)
                    if isinstance(obj, dict):
                        for k, v in obj.items():
                            if isinstance(v, list):
                                mapping[str(k).strip()] = [str(x).strip() for x in v if str(x).strip()]
                            elif isinstance(v, str):
                                mapping[str(k).strip()] = [e.strip() for e in v.split(',') if e.strip()]
                except Exception:
                    # fallback parse "Unit=mail1,mail2; Unit2=mail3"
                    parts = [p for p in raw.split(';') if p.strip()]
                    for p in parts:
                        if '=' in p:
                            name, emails = p.split('=', 1)
                            mapping[name.strip()] = [e.strip() for e in emails.split(',') if e.strip()]
            # Build form
            inputs = {}
            for u in units:
                emails = ", ".join(mapping.get(u.name, []))
                line = QLineEdit(emails)
                uf.addRow(u.name, line)
                inputs[u.name] = line
            # Buttons
            from PySide6.QtWidgets import QPushButton, QHBoxLayout
            ub_ok = QPushButton("Lưu"); ub_cancel = QPushButton("Đóng")
            urow = QHBoxLayout(); urow.addWidget(ub_ok); urow.addWidget(ub_cancel)
            uf.addRow(urow)

            def do_save_unit_emails():
                data = {}
                for name, line in inputs.items():
                    vals = [e.strip() for e in (line.text() or '').split(',') if e.strip()]
                    if vals:
                        data[name] = vals
                try:
                    set_setting('UNIT_EMAILS', json.dumps(data, ensure_ascii=False))
                    QMessageBox.information(udlg, "Đã lưu", "Đã lưu email theo đơn vị")
                except Exception as ex:
                    QMessageBox.critical(udlg, "Lỗi", str(ex))

            ub_ok.clicked.connect(do_save_unit_emails)
            ub_cancel.clicked.connect(udlg.reject)
            udlg.exec()

        def manage_unit_emails_db():
            # CRUD đơn giản cho bảng unit_email_recipients
            from PySide6.QtWidgets import QDialog, QVBoxLayout, QHBoxLayout, QTableWidget, QTableWidgetItem, QPushButton, QLineEdit, QComboBox
            md = QDialog(dlg); md.setWindowTitle("Quản lý email theo đơn vị (DB)")
            lay = QVBoxLayout(md)
            top = QHBoxLayout();
            unit_box = QComboBox();
            from .db import SessionLocal as _SL
            from .models import Unit, UnitEmailRecipient
            s = _SL(); units = s.query(Unit).order_by(Unit.name).all(); s.close()
            unit_map = {}
            unit_box.addItem("-- Chọn đơn vị --", None)
            for u in units:
                unit_box.addItem(u.name, u.id); unit_map[u.id] = u.name
            btn_load = QPushButton("Tải"); btn_add = QPushButton("Thêm"); btn_del = QPushButton("Xoá"); btn_toggle = QPushButton("Bật/Tắt"); btn_save_note = QPushButton("Lưu Note")
            top.addWidget(unit_box); top.addWidget(btn_load); top.addWidget(btn_add); top.addWidget(btn_del); top.addWidget(btn_toggle); top.addWidget(btn_save_note)
            lay.addLayout(top)
            table = QTableWidget(0, 4); table.setHorizontalHeaderLabels(["Email", "Active", "Note", "Created"])
            table.setEditTriggers(QTableWidget.DoubleClicked | QTableWidget.SelectedClicked)
            lay.addWidget(table)
            email_edit = QLineEdit(); email_note = QLineEdit();
            from PySide6.QtWidgets import QFormLayout as _QF
            sf = _QF(); sf.addRow("Email(s)", email_edit); sf.addRow("Note", email_note); lay.addLayout(sf)
            # Import từ settings
            btn_import_settings = QPushButton("Nhập từ settings")
            lay.addWidget(btn_import_settings)
            guard = {'updating': False}
            def load_rows():
                guard['updating'] = True
                table.setRowCount(0)
                uid = unit_box.currentData();
                if not uid:
                    guard['updating'] = False
                    return
                s2 = _SL()
                try:
                    rows = s2.query(UnitEmailRecipient).filter(UnitEmailRecipient.unit_id==uid).order_by(UnitEmailRecipient.created_at.desc()).all()
                    for r in rows:
                        i = table.rowCount(); table.insertRow(i)
                        table.setItem(i, 0, QTableWidgetItem(r.email or ''))
                        table.setItem(i, 1, QTableWidgetItem('1' if r.active else '0'))
                        table.setItem(i, 2, QTableWidgetItem(r.note or ''))
                        table.setItem(i, 3, QTableWidgetItem(str(r.created_at)))
                finally:
                    s2.close()
                guard['updating'] = False
            def do_add():
                uid = unit_box.currentData();
                em_raw = (email_edit.text() or '').strip()
                if not uid or not em_raw:
                    return
                # Cho phép nhập nhiều email, phân tách bởi dấu phẩy
                emails = [e.strip() for e in em_raw.split(',') if e.strip()]
                # Validate định dạng đơn giản
                import re
                pat = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")
                valid = [e for e in emails if pat.match(e)]
                invalid = [e for e in emails if e not in valid]
                if invalid:
                    QMessageBox.warning(md, "Cảnh báo", f"Bỏ qua email không hợp lệ: {', '.join(invalid)}")
                if not valid:
                    return
                s2 = _SL();
                try:
                    # tránh trùng: kiểm tra trước khi thêm
                    from sqlalchemy import exists
                    added = 0
                    note_text = (email_note.text() or '').strip()
                    for em in valid:
                        exists_q = s2.query(UnitEmailRecipient).filter(UnitEmailRecipient.unit_id==uid, UnitEmailRecipient.email==em).first()
                        if exists_q:
                            continue
                        s2.add(UnitEmailRecipient(unit_id=uid, email=em, active=True, note=note_text)); s2.commit(); added += 1
                    if added:
                        QMessageBox.information(md, "Đã thêm", f"Đã thêm {added} email")
                except Exception as ex:
                    s2.rollback(); QMessageBox.critical(md, "Lỗi", str(ex))
                finally:
                    s2.close(); load_rows()
            def do_del():
                uid = unit_box.currentData();
                row = table.currentRow()
                if not uid or row < 0:
                    return
                em = table.item(row,0).text() if table.item(row,0) else ''
                s2 = _SL();
                try:
                    s2.query(UnitEmailRecipient).filter(UnitEmailRecipient.unit_id==uid, UnitEmailRecipient.email==em).delete(); s2.commit()
                except Exception as ex:
                    s2.rollback(); QMessageBox.critical(md, "Lỗi", str(ex))
                finally:
                    s2.close(); load_rows()
            def do_toggle():
                uid = unit_box.currentData();
                row = table.currentRow()
                if not uid or row < 0:
                    return
                em = table.item(row,0).text() if table.item(row,0) else ''
                s2 = _SL();
                try:
                    rec = s2.query(UnitEmailRecipient).filter(UnitEmailRecipient.unit_id==uid, UnitEmailRecipient.email==em).first()
                    if rec:
                        rec.active = not bool(rec.active); s2.commit()
                except Exception as ex:
                    s2.rollback(); QMessageBox.critical(md, "Lỗi", str(ex))
                finally:
                    s2.close(); load_rows()
            def do_save_note():
                uid = unit_box.currentData();
                row = table.currentRow()
                if not uid or row < 0:
                    return
                em = table.item(row,0).text() if table.item(row,0) else ''
                note_val = table.item(row,2).text() if table.item(row,2) else ''
                s2 = _SL();
                try:
                    rec = s2.query(UnitEmailRecipient).filter(UnitEmailRecipient.unit_id==uid, UnitEmailRecipient.email==em).first()
                    if rec:
                        rec.note = note_val; s2.commit()
                except Exception as ex:
                    s2.rollback(); QMessageBox.critical(md, "Lỗi", str(ex))
                finally:
                    s2.close(); load_rows()
            def on_item_changed(it):
                if guard['updating']:
                    return
                # Cột Note (index 2)
                if it.column() != 2:
                    return
                do_save_note()
            def do_import_settings():
                try:
                    from .settings_service import get_setting
                    raw = get_setting('UNIT_EMAILS','') or ''
                    if not raw.strip():
                        QMessageBox.information(md, "Không có dữ liệu", "UNIT_EMAILS đang trống")
                        return
                    import json
                    mapping: dict[str, list[str]] = {}
                    try:
                        obj = json.loads(raw)
                        if isinstance(obj, dict):
                            for k, v in obj.items():
                                if isinstance(v, list):
                                    emails = [str(x).strip() for x in v if str(x).strip()]
                                else:
                                    emails = [e.strip() for e in str(v).split(',') if e.strip()]
                                if emails:
                                    mapping[str(k).strip()] = emails
                    except Exception:
                        parts = [p for p in raw.split(';') if p.strip()]
                        for p in parts:
                            if '=' in p:
                                name, emails = p.split('=', 1)
                                vals = [e.strip() for e in emails.split(',') if e.strip()]
                                if vals:
                                    mapping[name.strip()] = vals
                    if not mapping:
                        QMessageBox.information(md, "Không có dữ liệu", "Không thể parse UNIT_EMAILS")
                        return
                    s2 = _SL();
                    try:
                        # duyệt theo tên đơn vị, map sang id
                        units_by_name = {u.name.strip().lower(): u.id for u in s2.query(Unit).all()}
                        added = 0
                        import re
                        pat = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")
                        for name, emails in mapping.items():
                            uid = units_by_name.get((name or '').strip().lower())
                            if not uid:
                                continue
                            for em in emails:
                                em = (em or '').strip()
                                if not em or not pat.match(em):
                                    continue
                                try:
                                    # tránh trùng
                                    if s2.query(UnitEmailRecipient).filter(UnitEmailRecipient.unit_id==uid, UnitEmailRecipient.email==em).first():
                                        continue
                                    s2.add(UnitEmailRecipient(unit_id=uid, email=em, active=True)); s2.commit(); added += 1
                                    # Audit
                                    try:
                                        log_action(SessionLocal(), self.current_user.get('id'), 'unit_email_import', 'UnitEmailRecipient', None, f"unit={name};email={em}")
                                    except Exception:
                                        pass
                                except Exception:
                                    s2.rollback()
                        QMessageBox.information(md, "Đã nhập", f"Đã thêm {added} địa chỉ")
                    finally:
                        s2.close(); load_rows()
                except Exception as ex:
                    QMessageBox.critical(md, "Lỗi", str(ex))
            table.itemChanged.connect(on_item_changed)
            btn_load.clicked.connect(load_rows); btn_add.clicked.connect(do_add); btn_del.clicked.connect(do_del); btn_toggle.clicked.connect(do_toggle); btn_save_note.clicked.connect(do_save_note); btn_import_settings.clicked.connect(do_import_settings)
        btn_ok.clicked.connect(save)
        btn_cancel.clicked.connect(dlg.reject)
        btn_unit_emails.clicked.connect(edit_unit_emails)
        try:
            btn_unit_emails_db.clicked.connect(manage_unit_emails_db)
        except Exception:
            pass
        def test_smtp():
            try:
                from .mailer import send_alert
                ok = send_alert('[HRMS] SMTP test', 'This is a test email from HRMS settings')
                if ok:
                    QMessageBox.information(dlg, "OK", "Đã gửi email thử nghiệm (kiểm tra hộp thư ALERT_EMAILS)")
                else:
                    QMessageBox.warning(dlg, "Chưa gửi", "Gửi email thử nghiệm thất bại — kiểm tra cấu hình SMTP/ALERT_EMAILS")
            except Exception as ex:
                QMessageBox.critical(dlg, "Lỗi", str(ex))
        btn_test_smtp.clicked.connect(test_smtp)
        dlg.exec()

    def open_email_history(self):
        # Hiển thị lịch sử email với bộ lọc cơ bản
        role = (self.current_user.get('role') or '').lower()
        if role not in ('admin','hr'):
            QMessageBox.warning(self, "Không có quyền", "Chỉ admin/HR mới truy cập")
            return
        from PySide6.QtWidgets import QDialog, QFormLayout, QTableWidget, QTableWidgetItem, QPushButton, QCheckBox, QFileDialog, QComboBox, QLabel, QHBoxLayout
        dlg = QDialog(self); dlg.setWindowTitle("Lịch sử Email")
        lay = QVBoxLayout(dlg)
        # Bộ lọc
        f = QFormLayout();
        from PySide6.QtWidgets import QLineEdit, QComboBox, QPushButton, QDateEdit, QCheckBox
        type_box = QComboBox(); type_box.addItems(["(Tất cả)", "salary_due", "bhxh_monthly", "contracts_expiring", "quick_report", "retirement", "generic"])
        unit_edit = QLineEdit(); unit_edit.setPlaceholderText("Đơn vị chứa…")
        unit_box = QComboBox(); unit_box.addItem("(Tất cả)")
        def reload_units():
            try:
                unit_box.blockSignals(True)
                cur = unit_box.currentText()
                unit_box.clear(); unit_box.addItem("(Tất cả)")
                from .db import SessionLocal as _SL
                from .models import Unit as _Unit
                _s = _SL();
                try:
                    for _u in _s.query(_Unit).order_by(_Unit.name).all():
                        unit_box.addItem(_u.name)
                finally:
                    _s.close()
                # cố gắng chọn lại mục cũ
                for i in range(unit_box.count()):
                    if unit_box.itemText(i) == cur:
                        unit_box.setCurrentIndex(i); break
            except Exception:
                pass
            finally:
                try: unit_box.blockSignals(False)
                except Exception: pass
        reload_units()
        status_box = QComboBox(); status_box.addItems(["(Tất cả)", "sent", "failed"])
        only_failed = QCheckBox("Chỉ lỗi")
        subject_search = QLineEdit(); subject_search.setPlaceholderText("Tìm tiêu đề…")
        from_date = QDateEdit(); from_date.setCalendarPopup(True)
        to_date = QDateEdit(); to_date.setCalendarPopup(True)
        e_from = QCheckBox("Áp dụng"); e_to = QCheckBox("Áp dụng")
        btn_refresh = QPushButton("Làm mới"); btn_export = QPushButton("Export CSV"); btn_view = QPushButton("Xem chi tiết"); btn_save_filter = QPushButton("Lưu bộ lọc"); btn_reset_filter = QPushButton("Reset bộ lọc"); btn_export_zip = QPushButton("Export ZIP")
        f.addRow("Loại", type_box)
        f.addRow("Đơn vị", unit_edit)
        from PySide6.QtWidgets import QHBoxLayout as _HB, QPushButton as _PB
        _unit_row = _HB(); _unit_row.addWidget(unit_box); _btn_reload_units = _PB("Tải lại")
        _btn_reload_units.clicked.connect(reload_units)
        _unit_row.addWidget(_btn_reload_units)
        f.addRow("Chọn đơn vị", _unit_row)
        f.addRow("Trạng thái", status_box)
        f.addRow("", only_failed)
        f.addRow("Tiêu đề", subject_search)
        h = QHBoxLayout(); h.addWidget(QLabel("Từ")); h.addWidget(from_date); h.addWidget(e_from); h.addSpacing(12); h.addWidget(QLabel("Đến")); h.addWidget(to_date); h.addWidget(e_to)
        f.addRow("Thời gian", h)
        user_id_edit = QLineEdit(); user_id_edit.setPlaceholderText("User ID")
        f.addRow("User ID", user_id_edit)
        # Phân trang
        pager_bar = QHBoxLayout()
        page_size_box = QComboBox(); page_size_box.addItems(["100","200","500"]) ; page_size_box.setCurrentText("100")
        btn_prev = QPushButton("◀ Trước"); btn_next = QPushButton("Sau ▶"); page_label = QLabel("Trang 1/1")
        from PySide6.QtWidgets import QLineEdit as _QLE
        page_input = _QLE(); page_input.setPlaceholderText("Trang..."); btn_go = QPushButton("Đi")
        pager_bar.addWidget(QLabel("Kích thước trang")); pager_bar.addWidget(page_size_box); pager_bar.addWidget(btn_prev); pager_bar.addWidget(btn_next); pager_bar.addWidget(page_label); pager_bar.addWidget(page_input); pager_bar.addWidget(btn_go)
        f.addRow(pager_bar)
        # Áp dụng trên toàn bộ (bỏ qua phân trang)
        cb_all_scope = QCheckBox("Toàn bộ (bỏ qua phân trang)")
        f.addRow("", cb_all_scope)
        btn_resend = QPushButton("Gửi lại")
        btn_resend_all = QPushButton("Gửi lại tất cả (lọc)")
        btn_resend_group = QPushButton("Gửi lại theo nhóm (đơn vị)")
        btn_view_zip = QPushButton("Xem ZIP")
        btn_open_files = QPushButton("Xem file")
        btn_open_folders = QPushButton("Mở thư mục file")
        btn_open_exports = QPushButton("Mở exports")
        btn_copy_recip = QPushButton("Copy recipients")
        btn_delete = QPushButton("Xoá")
        btn_delete_all = QPushButton("Xoá tất cả (lọc)")
        hb = QHBoxLayout();
        for b in (btn_refresh, btn_save_filter, btn_reset_filter, btn_export, btn_export_zip, btn_view, btn_resend, btn_resend_all, btn_resend_group, btn_view_zip, btn_open_files, btn_open_folders, btn_open_exports, btn_copy_recip, btn_delete, btn_delete_all):
            hb.addWidget(b)
        f.addRow(hb)
        lay.addLayout(f)
        # Bảng kết quả
        table = QTableWidget(0, 8)
        table.setHorizontalHeaderLabels(["Thời gian", "Loại", "Đơn vị", "Subject", "Recipients", "Tệp đính kèm", "Trạng thái", "Lỗi"])
        lay.addWidget(table)
        # Trạng thái phân trang
        state = {'page': 0, 'total': 0}
        # Khôi phục bộ lọc theo user
        def load_filter():
            try:
                from .settings_service import get_setting
                key = f"EMAIL_HISTORY_FILTER:{(self.current_user.get('username') or '').strip()}"
                raw = get_setting(key, '') or ''
                if not raw: return
                import json
                obj = json.loads(raw)
                # type
                t = obj.get('type')
                if t and t in [type_box.itemText(i) for i in range(type_box.count())]:
                    type_box.setCurrentText(t)
                # unit exact
                usel = obj.get('unit_selected')
                if usel:
                    unit_box.setCurrentText(usel)
                # unit contains
                unit_edit.setText(obj.get('unit_contains',''))
                # status
                st = obj.get('status')
                if st and st in [status_box.itemText(i) for i in range(status_box.count())]:
                    status_box.setCurrentText(st)
                only_failed.setChecked(bool(obj.get('only_failed', False)))
                subject_search.setText(obj.get('subject',''))
                # dates
                from PySide6.QtCore import QDate
                if obj.get('from_date'):
                    y,m,d = obj['from_date'].split('-')
                    from_date.setDate(QDate(int(y), int(m), int(d)))
                    e_from.setChecked(True)
                if obj.get('to_date'):
                    y,m,d = obj['to_date'].split('-')
                    to_date.setDate(QDate(int(y), int(m), int(d)))
                    e_to.setChecked(True)
                uid = obj.get('user_id')
                user_id_edit.setText(str(uid) if uid is not None else '')
                # group resend fields
                try:
                    if 'zip_pattern' in obj and hasattr(resend_grouped_by_unit, '__zip_pattern'):
                        resend_grouped_by_unit.__zip_pattern.setText(obj.get('zip_pattern') or resend_grouped_by_unit.__zip_pattern.text())
                    if 'subj_prefix' in obj and hasattr(resend_grouped_by_unit, '__subj_prefix'):
                        resend_grouped_by_unit.__subj_prefix.setText(obj.get('subj_prefix') or '')
                    if 'subj_suffix' in obj and hasattr(resend_grouped_by_unit, '__subj_suffix'):
                        resend_grouped_by_unit.__subj_suffix.setText(obj.get('subj_suffix') or '')
                except Exception:
                    pass
                # all scope
                try:
                    cb_all_scope.setChecked(bool(obj.get('all_scope', False)))
                except Exception:
                    pass
                # group zip checkbox
                try:
                    if hasattr(resend_grouped_by_unit, '__zip_cb'):
                        resend_grouped_by_unit.__zip_cb.setChecked(bool(obj.get('zip_when_group', False)))
                except Exception:
                    pass
                # page size
                try:
                    ps = str(obj.get('page_size','') or '')
                    if ps and ps in [page_size_box.itemText(i) for i in range(page_size_box.count())]:
                        page_size_box.setCurrentText(ps)
                except Exception:
                    pass
            except Exception:
                pass
        def save_filter():
            try:
                from .settings_service import set_setting
                key = f"EMAIL_HISTORY_FILTER:{(self.current_user.get('username') or '').strip()}"
                from datetime import date as _date
                def qdate_to_str(qd):
                    try:
                        return f"{qd.year()}-{qd.month():02d}-{qd.day():02d}"
                    except Exception:
                        return None
                obj = {
                    'type': type_box.currentText(),
                    'unit_selected': unit_box.currentText() if unit_box.currentIndex()>0 else '',
                    'unit_contains': unit_edit.text().strip(),
                    'status': status_box.currentText(),
                    'only_failed': only_failed.isChecked(),
                    'subject': subject_search.text().strip(),
                    'from_date': qdate_to_str(from_date.date()) if e_from.isChecked() else None,
                    'to_date': qdate_to_str(to_date.date()) if e_to.isChecked() else None,
                    'user_id': (int(user_id_edit.text()) if user_id_edit.text().isdigit() else None),
                    'zip_pattern': getattr(resend_grouped_by_unit, '__zip_pattern').text() if hasattr(resend_grouped_by_unit, '__zip_pattern') else None,
                    'subj_prefix': getattr(resend_grouped_by_unit, '__subj_prefix').text() if hasattr(resend_grouped_by_unit, '__subj_prefix') else None,
                    'subj_suffix': getattr(resend_grouped_by_unit, '__subj_suffix').text() if hasattr(resend_grouped_by_unit, '__subj_suffix') else None,
                    'all_scope': cb_all_scope.isChecked(),
                    'zip_when_group': (getattr(resend_grouped_by_unit, '__zip_cb').isChecked() if hasattr(resend_grouped_by_unit, '__zip_cb') else False),
                    'page_size': page_size_box.currentText(),
                }
                import json
                set_setting(key, json.dumps(obj, ensure_ascii=False))
            except Exception:
                pass
        load_filter()
        # Tải dữ liệu
        def load():
            from .db import SessionLocal
            from .models import EmailLog
            db = SessionLocal()
            try:
                q = db.query(EmailLog).order_by(EmailLog.created_at.desc())
                t = type_box.currentText();
                if not t.startswith("("):
                    q = q.filter(EmailLog.type == t)
                u_sel = unit_box.currentText().strip() if unit_box.currentIndex() > -1 else ""
                if u_sel and not u_sel.startswith("("):
                    q = q.filter(EmailLog.unit_name == u_sel)
                else:
                    u = unit_edit.text().strip();
                    if u:
                        q = q.filter(EmailLog.unit_name.ilike(f"%{u}%"))
                st = status_box.currentText();
                if only_failed.isChecked():
                    q = q.filter(EmailLog.status == 'failed')
                elif not st.startswith("("):
                    q = q.filter(EmailLog.status == st)
                subj = subject_search.text().strip();
                if subj:
                    q = q.filter(EmailLog.subject.ilike(f"%{subj}%"))
                uidtxt = user_id_edit.text().strip();
                if uidtxt.isdigit():
                    q = q.filter(EmailLog.user_id == int(uidtxt))
                # Lọc theo thời gian nếu bật
                if e_from.isChecked():
                    fd = from_date.date()
                    from datetime import datetime as _dt
                    q = q.filter(EmailLog.created_at >= _dt(fd.year(), fd.month(), fd.day(), 0, 0, 0))
                if e_to.isChecked():
                    td = to_date.date()
                    from datetime import datetime as _dt
                    q = q.filter(EmailLog.created_at <= _dt(td.year(), td.month(), td.day(), 23, 59, 59))
                # Áp dụng phân trang
                page_size = int(page_size_box.currentText()) if page_size_box.currentText().isdigit() else 100
                state['total'] = q.count()
                max_page = max(1, (state['total'] - 1)//page_size + 1)
                if state['page'] >= max_page:
                    state['page'] = max_page - 1
                if state['page'] < 0:
                    state['page'] = 0
                offset = state['page'] * page_size
                rows = q.order_by(EmailLog.created_at.desc()).limit(page_size).offset(offset).all()
                table.setRowCount(0)
                for r in rows:
                    i = table.rowCount(); table.insertRow(i)
                    it0 = QTableWidgetItem(str(getattr(r, 'created_at', '')))
                    it0.setData(Qt.UserRole, getattr(r, 'id', None))
                    it1 = QTableWidgetItem(getattr(r, 'type', '') or '')
                    it2 = QTableWidgetItem(getattr(r, 'unit_name', '') or '')
                    it3 = QTableWidgetItem(getattr(r, 'subject', '') or '')
                    it4 = QTableWidgetItem((getattr(r, 'recipients', '') or '')[:100])
                    it5 = QTableWidgetItem((getattr(r, 'attachments', '') or '')[:100])
                    it6 = QTableWidgetItem(getattr(r, 'status', '') or '')
                    it7 = QTableWidgetItem(getattr(r, 'error', '') or '')
                    # Lưu full text ở UserRole để xem chi tiết
                    it3.setData(Qt.UserRole, getattr(r, 'body', '') or '')
                    it4.setData(Qt.UserRole, getattr(r, 'recipients', '') or '')
                    it5.setData(Qt.UserRole, getattr(r, 'attachments', '') or '')
                    it7.setData(Qt.UserRole, getattr(r, 'error', '') or '')
                    table.setItem(i, 0, it0)
                    table.setItem(i, 1, it1)
                    table.setItem(i, 2, it2)
                    table.setItem(i, 3, it3)
                    table.setItem(i, 4, it4)
                    table.setItem(i, 5, it5)
                    table.setItem(i, 6, it6)
                    table.setItem(i, 7, it7)
                # Cập nhật pager
                page_label.setText(f"Trang {state['page']+1}/{max_page} ({state['total']} kết quả)")
                btn_prev.setEnabled(state['page']>0)
                btn_next.setEnabled(state['page']<max_page-1)
            except Exception as ex:
                QMessageBox.critical(dlg, "Lỗi", str(ex))
            finally:
                db.close()
        def export_csv():
            try:
                import csv
                from datetime import datetime as _dt
                from pathlib import Path
                Path('exports').mkdir(exist_ok=True)
                outp = Path('exports')/f"email_logs_{_dt.now().strftime('%Y%m%d_%H%M%S')}.csv"
                with open(outp, 'w', encoding='utf-8', newline='') as fcsv:
                    w = csv.writer(fcsv)
                    w.writerow(["created_at","type","unit_name","subject","recipients","attachments","status","error"])
                    for i in range(table.rowCount()):
                        it = lambda r,c: (table.item(r,c).text() if table.item(r,c) else '')
                        w.writerow([it(i,0), it(i,1), it(i,2), it(i,3), it(i,4), it(i,5), it(i,6), it(i,7)])
                QMessageBox.information(dlg, "Đã xuất", f"{outp}")
            except Exception as ex:
                QMessageBox.critical(dlg, "Lỗi", str(ex))
        def go_prev():
            state['page'] = max(0, state['page'] - 1)
            load()
        def go_next():
            # sẽ được giới hạn trong load
            state['page'] += 1
            load()
        btn_prev.clicked.connect(go_prev)
        btn_next.clicked.connect(go_next)
        def go_to_page():
            try:
                val = int(page_input.text()) if page_input.text().isdigit() else 1
                val = max(1, val)
                # Clamp to max page
                page_size = int(page_size_box.currentText()) if page_size_box.currentText().isdigit() else 100
                max_page = max(1, (state.get('total', 0) - 1)//page_size + 1)
                if val > max_page:
                    from PySide6.QtWidgets import QMessageBox as _QMB
                    _QMB.information(dlg, "Vượt phạm vi", f"Số trang tối đa hiện tại: {max_page}")
                    val = max_page
                state['page'] = val - 1
                load()
            except Exception:
                pass
        btn_go.clicked.connect(go_to_page)
        page_size_box.currentTextChanged.connect(lambda _ : (state.__setitem__('page', 0), load()))
        btn_refresh.clicked.connect(lambda: (save_filter(), state.__setitem__('page', 0), load()))
        btn_save_filter.clicked.connect(save_filter)
        def reset_filter():
            try:
                # Clear saved filter and UI fields
                from .settings_service import set_setting
                key = f"EMAIL_HISTORY_FILTER:{(self.current_user.get('username') or '').strip()}"
                set_setting(key, '')
            except Exception:
                pass
            try:
                type_box.setCurrentIndex(0)
                unit_box.setCurrentIndex(0)
                unit_edit.clear()
                status_box.setCurrentIndex(0)
                only_failed.setChecked(False)
                subject_search.clear()
                e_from.setChecked(False)
                e_to.setChecked(False)
                user_id_edit.clear()
                state['page'] = 0
                load()
            except Exception:
                pass
        btn_reset_filter.clicked.connect(reset_filter)
        btn_export.clicked.connect(export_csv)
        def export_failed_csv():
            try:
                import csv
                from datetime import datetime as _dt
                from pathlib import Path
                from .db import SessionLocal
                from .models import EmailLog
                Path('exports').mkdir(exist_ok=True)
                outp = Path('exports')/f"email_logs_failed_{_dt.now().strftime('%Y%m%d_%H%M%S')}.csv"
                dbf = SessionLocal()
                try:
                    q = dbf.query(EmailLog).filter(EmailLog.status=='failed')
                    # reuse basic filters except status
                    t = type_box.currentText()
                    if not t.startswith("("):
                        q = q.filter(EmailLog.type == t)
                    u_sel = unit_box.currentText().strip() if unit_box.currentIndex() > -1 else ""
                    if u_sel and not u_sel.startswith("("):
                        q = q.filter(EmailLog.unit_name == u_sel)
                    else:
                        u = unit_edit.text().strip()
                        if u:
                            q = q.filter(EmailLog.unit_name.ilike(f"%{u}%"))
                    subj = subject_search.text().strip()
                    if subj:
                        q = q.filter(EmailLog.subject.ilike(f"%{subj}%"))
                    uidtxt = user_id_edit.text().strip()
                    if uidtxt.isdigit():
                        q = q.filter(EmailLog.user_id == int(uidtxt))
                    rows = q.order_by(EmailLog.created_at.desc()).all()
                    with open(outp, 'w', encoding='utf-8', newline='') as fcsv:
                        w = csv.writer(fcsv)
                        w.writerow(["created_at","type","unit_name","subject","recipients","attachments","status","error"])
                        for r in rows:
                            w.writerow([str(getattr(r,'created_at','')), getattr(r,'type','') or '', getattr(r,'unit_name','') or '', getattr(r,'subject','') or '', getattr(r,'recipients','') or '', getattr(r,'attachments','') or '', getattr(r,'status','') or '', getattr(r,'error','') or ''])
                    QMessageBox.information(dlg, "Đã xuất", f"{outp}")
                finally:
                    dbf.close()
            except Exception as ex:
                QMessageBox.critical(dlg, "Lỗi", str(ex))
        btn_export_failed = QPushButton("Export lỗi CSV"); btn_export_failed.clicked.connect(export_failed_csv); f.addRow(btn_export_failed)
        def export_attachments_zip():
            try:
                from .db import SessionLocal
                from .models import EmailLog
                import zipfile, os
                from pathlib import Path
                from datetime import datetime as _dt
                db2 = SessionLocal()
                try:
                    q = db2.query(EmailLog)
                    t = type_box.currentText()
                    if not t.startswith("("):
                        q = q.filter(EmailLog.type == t)
                    u_sel = unit_box.currentText().strip() if unit_box.currentIndex() > -1 else ""
                    if u_sel and not u_sel.startswith("("):
                        q = q.filter(EmailLog.unit_name == u_sel)
                    else:
                        u = unit_edit.text().strip()
                        if u:
                            q = q.filter(EmailLog.unit_name.ilike(f"%{u}%"))
                    st = status_box.currentText()
                    if only_failed.isChecked():
                        q = q.filter(EmailLog.status == 'failed')
                    elif not st.startswith("("):
                        q = q.filter(EmailLog.status == st)
                    subj = subject_search.text().strip()
                    if subj:
                        q = q.filter(EmailLog.subject.ilike(f"%{subj}%"))
                    uidtxt = user_id_edit.text().strip()
                    if uidtxt.isdigit():
                        q = q.filter(EmailLog.user_id == int(uidtxt))
                    # Date range
                    from datetime import datetime as _dt2
                    if e_from.isChecked():
                        fd = from_date.date(); q = q.filter(EmailLog.created_at >= _dt2(fd.year(), fd.month(), fd.day(), 0, 0, 0))
                    if e_to.isChecked():
                        td = to_date.date(); q = q.filter(EmailLog.created_at <= _dt2(td.year(), td.month(), td.day(), 23, 59, 59))
                    rows = q.order_by(EmailLog.created_at.desc()).all()
                    files = []
                    total_size = 0
                    for r in rows:
                        for part in (getattr(r,'attachments','') or '').split(','):
                            p = part.strip()
                            if not p:
                                continue
                            fp = Path(p)
                            if fp.exists():
                                try:
                                    total_size += fp.stat().st_size
                                except Exception:
                                    pass
                                files.append((getattr(r,'unit_name','') or '', getattr(r,'type','') or 'generic', fp))
                    if not files:
                        QMessageBox.information(dlg, "Không có file", "Không có tệp đính kèm hợp lệ trong kết quả lọc")
                        return
                    # Xác nhận trước khi gom ZIP
                    def _hsz(n):
                        try:
                            for unit in ['B','KB','MB','GB','TB']:
                                if n < 1024.0:
                                    return f"{n:0.1f} {unit}"
                                n /= 1024.0
                            return f"{n:0.1f} PB"
                        except Exception:
                            return str(n)
                    from PySide6.QtWidgets import QMessageBox as _QMB
                    if _QMB.question(dlg, "Xác nhận", f"Sẽ gom {len(files)} tệp (~{_hsz(total_size)}). Tiếp tục?", _QMB.Yes|_QMB.No) != _QMB.Yes:
                        return
                    Path('exports').mkdir(exist_ok=True)
                    ts = _dt.now().strftime('%Y%m%d_%H%M%S')
                    zip_path = Path('exports')/f"email_attachments_{ts}.zip"
                    with zipfile.ZipFile(str(zip_path), 'w', compression=zipfile.ZIP_DEFLATED) as zf:
                        added = set()
                        for unit_name, etype, fp in files:
                            arcdir = f"{(unit_name or 'NoUnit').replace(' ','_')}/{(etype or 'generic')}"
                            arcname = f"{arcdir}/{fp.name}"
                            key = (arcname.lower(), fp.resolve())
                            if key in added:
                                continue
                            try:
                                zf.write(str(fp), arcname=arcname)
                                added.add(key)
                            except Exception:
                                continue
                    QMessageBox.information(dlg, "Đã xuất", f"{zip_path}")
                    # Audit + Log export
                    try:
                        from .db import SessionLocal as _SL
                        from .models import EmailLog as _EL
                        # Audit
                        try:
                            log_action(SessionLocal(), self.current_user.get('id'), 'export_email_attachments_zip', 'EmailLog', None, f"zip={zip_path};files={len(files)};size={total_size}")
                        except Exception:
                            pass
                        s = _SL(); s.add(_EL(type='export_zip', unit_name=None, recipients='', subject='Export attachments ZIP', body='Filtered attachments ZIP', attachments=str(zip_path), status='exported', user_id=self.current_user.get('id'))); s.commit(); s.close()
                    except Exception:
                        pass
                finally:
                    db2.close()
            except Exception as ex:
                QMessageBox.critical(dlg, "Lỗi", str(ex))
        btn_export_zip.clicked.connect(export_attachments_zip)
        def view_detail():
            i = table.currentRow()
            if i < 0:
                QMessageBox.information(dlg, "Chưa chọn", "Chọn một dòng trong bảng")
                return
            from PySide6.QtWidgets import QDialog, QFormLayout, QTextEdit
            dd = QDialog(dlg); dd.setWindowTitle("Chi tiết Email")
            ff = QFormLayout(dd)
            get = lambda c: (table.item(i,c).data(Qt.UserRole) if table.item(i,c) else '') or (table.item(i,c).text() if table.item(i,c) else '')
            ff.addRow("Thời gian", QLabel(table.item(i,0).text() if table.item(i,0) else ''))
            ff.addRow("Loại", QLabel(table.item(i,1).text() if table.item(i,1) else ''))
            ff.addRow("Đơn vị", QLabel(table.item(i,2).text() if table.item(i,2) else ''))
            ff.addRow("Subject", QLabel(table.item(i,3).text() if table.item(i,3) else ''))
            body = QTextEdit(); body.setReadOnly(True); body.setText(get(3))
            ff.addRow("Body", body)
            ff.addRow("Recipients", QLabel(get(4)))
            ff.addRow("Attachments", QLabel(get(5)))
            ff.addRow("Trạng thái", QLabel(table.item(i,6).text() if table.item(i,6) else ''))
            ff.addRow("Lỗi", QLabel(get(7)))
            dd.resize(700, 500)
            dd.exec()
        btn_view.clicked.connect(view_detail)
        def resend_selected():
            i = table.currentRow()
            if i < 0:
                QMessageBox.information(dlg, "Chưa chọn", "Chọn một dòng để gửi lại")
                return
            try:
                t = table.item(i,1).text() if table.item(i,1) else ''
                unit_name = table.item(i,2).text() if table.item(i,2) else ''
                subject = table.item(i,3).text() if table.item(i,3) else ''
                body = table.item(i,3).data(Qt.UserRole) if table.item(i,3) else ''
                attach_disp = table.item(i,5).data(Qt.UserRole) if table.item(i,5) else (table.item(i,5).text() if table.item(i,5) else '')
                paths = []
                from pathlib import Path
                for part in (attach_disp or '').split(','):
                    p = part.strip()
                    if p:
                        pp = Path(p)
                        if pp.exists():
                            paths.append(str(pp))
                recips = None
                if t in ('salary_due','bhxh_monthly','contracts_expiring') and unit_name:
                    try:
                        from .mailer import get_recipients_for_unit
                        recips = get_recipients_for_unit(unit_name)
                    except Exception:
                        recips = None
                from .mailer import send_email_with_attachment
                ok = send_email_with_attachment(subject or '[HRMS] Resend', body or '', paths, to=recips)
                # Log lại với type ban đầu
                try:
                    from .db import SessionLocal
                    from .models import EmailLog
                    s = SessionLocal(); s.add(EmailLog(type=t or 'generic', unit_name=(unit_name or None), recipients='', subject=subject or '[HRMS] Resend', body=(body or '')[:1000], attachments=', '.join(paths), status='sent' if ok else 'failed', user_id=self.current_user.get('id'))); s.commit(); s.close()
                except Exception:
                    pass
                if ok:
                    QMessageBox.information(dlg, "Đã gửi lại", "Gửi lại thành công")
                else:
                    QMessageBox.warning(dlg, "Chưa gửi", "Gửi lại thất bại")
            except Exception as ex:
                QMessageBox.critical(dlg, "Lỗi", str(ex))
        btn_resend.clicked.connect(resend_selected)
        def open_selected_files():
            try:
                i = table.currentRow()
                if i < 0:
                    QMessageBox.information(dlg, "Chưa chọn", "Chọn một dòng để mở file")
                    return
                import os, sys, subprocess
                from pathlib import Path
                attach_disp = table.item(i,5).data(Qt.UserRole) if table.item(i,5) else (table.item(i,5).text() if table.item(i,5) else '')
                files = []
                for part in (attach_disp or '').split(','):
                    p = part.strip()
                    if p and Path(p).exists():
                        files.append(str(Path(p)))
                if not files:
                    QMessageBox.information(dlg, "Không có file", "Không có file hợp lệ để mở")
                    return
                for p in files:
                    if sys.platform.startswith('win'):
                        os.startfile(p)
                    elif sys.platform == 'darwin':
                        subprocess.Popen(['open', p])
                    else:
                        subprocess.Popen(['xdg-open', p])
            except Exception as ex:
                QMessageBox.critical(dlg, "Lỗi", str(ex))
        btn_open_files.clicked.connect(open_selected_files)
        def view_zip_contents():
            try:
                i = table.currentRow()
                if i < 0:
                    QMessageBox.information(dlg, "Chưa chọn", "Chọn một dòng để xem ZIP")
                    return
                from pathlib import Path
                attach_disp = table.item(i,5).data(Qt.UserRole) if table.item(i,5) else (table.item(i,5).text() if table.item(i,5) else '')
                zips = []
                for part in (attach_disp or '').split(','):
                    p = part.strip()
                    if p and p.lower().endswith('.zip') and Path(p).exists():
                        zips.append(Path(p))
                if not zips:
                    QMessageBox.information(dlg, "Không có ZIP", "Không có file ZIP hợp lệ trong đính kèm")
                    return
                # Hiển thị nội dung của ZIP đầu tiên
                import zipfile
                with zipfile.ZipFile(str(zips[0]), 'r') as zf:
                    infos = zf.infolist()
                from PySide6.QtWidgets import QDialog as _QD, QVBoxLayout as _QV, QTableWidget as _QT, QTableWidgetItem as _QTI, QPushButton as _QP, QHBoxLayout as _QH
                d = _QD(dlg); d.setWindowTitle(f"Nội dung ZIP: {zips[0].name}")
                lay2 = _QV(d)
                t = _QT(0, 3); t.setHorizontalHeaderLabels(["Tên", "Kích thước", "Nén"]) ; lay2.addWidget(t)
                for inf in infos:
                    j = t.rowCount(); t.insertRow(j)
                    t.setItem(j, 0, _QTI(inf.filename))
                    t.setItem(j, 1, _QTI(str(inf.file_size)))
                    t.setItem(j, 2, _QTI(str(inf.compress_size)))
                # Nút thao tác
                bar = _QH(); btn_open = _QP("Mở mục đã chọn"); btn_extract_all = _QP("Giải nén tất cả")
                bar.addWidget(btn_open); bar.addWidget(btn_extract_all); lay2.addLayout(bar)
                def do_open_selected():
                    try:
                        row = t.currentRow()
                        if row < 0:
                            return
                        name = t.item(row,0).text() if t.item(row,0) else ''
                        if not name:
                            return
                        # extract to temp and open
                        import os, sys, subprocess
                        import zipfile as _zf
                        out_dir = Path('temp')/'zipview'/zips[0].stem
                        out_dir.mkdir(parents=True, exist_ok=True)
                        with _zf.ZipFile(str(zips[0]), 'r') as _zz:
                            _zz.extract(name, str(out_dir))
                        target = out_dir/name
                        if sys.platform.startswith('win'):
                            os.startfile(str(target))
                        elif sys.platform == 'darwin':
                            subprocess.Popen(['open', str(target)])
                        else:
                            subprocess.Popen(['xdg-open', str(target)])
                    except Exception as ex2:
                        QMessageBox.critical(d, "Lỗi", str(ex2))
                def do_extract_all():
                    try:
                        import zipfile as _zf
                        out_dir = Path('temp')/'zipview'/zips[0].stem
                        out_dir.mkdir(parents=True, exist_ok=True)
                        with _zf.ZipFile(str(zips[0]), 'r') as _zz:
                            _zz.extractall(str(out_dir))
                        QMessageBox.information(d, "Đã giải nén", str(out_dir))
                    except Exception as ex3:
                        QMessageBox.critical(d, "Lỗi", str(ex3))
                btn_open.clicked.connect(do_open_selected)
                btn_extract_all.clicked.connect(do_extract_all)
                d.resize(700, 450); d.exec()
            except Exception as ex:
                QMessageBox.critical(dlg, "Lỗi", str(ex))
        btn_view_zip.clicked.connect(view_zip_contents)
        def open_selected_folders():
            try:
                i = table.currentRow()
                if i < 0:
                    QMessageBox.information(dlg, "Chưa chọn", "Chọn một dòng để mở thư mục")
                    return
                import os, sys, subprocess
                from pathlib import Path
                attach_disp = table.item(i,5).data(Qt.UserRole) if table.item(i,5) else (table.item(i,5).text() if table.item(i,5) else '')
                folders = set()
                for part in (attach_disp or '').split(','):
                    p = part.strip()
                    if p:
                        fp = Path(p)
                        if fp.exists():
                            folders.add(str(fp.parent))
                if not folders:
                    QMessageBox.information(dlg, "Không có thư mục", "Không có thư mục hợp lệ để mở")
                    return
                for d in folders:
                    if sys.platform.startswith('win'):
                        os.startfile(d)
                    elif sys.platform == 'darwin':
                        subprocess.Popen(['open', d])
                    else:
                        subprocess.Popen(['xdg-open', d])
            except Exception as ex:
                QMessageBox.critical(dlg, "Lỗi", str(ex))
        btn_open_folders.clicked.connect(open_selected_folders)
        def copy_recipients():
            try:
                i = table.currentRow()
                if i < 0:
                    QMessageBox.information(dlg, "Chưa chọn", "Chọn một dòng để copy recipients")
                    return
                from PySide6.QtWidgets import QApplication
                txt = table.item(i,4).data(Qt.UserRole) if table.item(i,4) else (table.item(i,4).text() if table.item(i,4) else '')
                QApplication.clipboard().setText(txt or '')
                QMessageBox.information(dlg, "Đã copy", "Recipients đã được copy vào clipboard")
            except Exception as ex:
                QMessageBox.critical(dlg, "Lỗi", str(ex))
        btn_copy_recip.clicked.connect(copy_recipients)
        def open_exports():
            try:
                import os, sys, subprocess
                p = 'exports'
                if sys.platform.startswith('win'):
                    os.startfile(p)
                elif sys.platform == 'darwin':
                    subprocess.Popen(['open', p])
                else:
                    subprocess.Popen(['xdg-open', p])
            except Exception as ex:
                QMessageBox.critical(dlg, "Lỗi", str(ex))
        btn_open_exports.clicked.connect(open_exports)
        def delete_selected():
            i = table.currentRow()
            if i < 0:
                QMessageBox.information(dlg, "Chưa chọn", "Chọn một dòng để xoá")
                return
            try:
                eid = table.item(i,0).data(Qt.UserRole) if table.item(i,0) else None
                if not eid:
                    QMessageBox.warning(dlg, "Không xác định", "Không xác định ID")
                    return
                from .db import SessionLocal
                from .models import EmailLog
                s = SessionLocal(); s.query(EmailLog).filter(EmailLog.id==eid).delete(); s.commit(); s.close()
                table.removeRow(i)
            except Exception as ex:
                QMessageBox.critical(dlg, "Lỗi", str(ex))
        btn_delete.clicked.connect(delete_selected)
        def delete_all_filtered():
            from PySide6.QtWidgets import QMessageBox as _QMB
            if _QMB.question(dlg, "Xác nhận", "Xoá tất cả bản ghi đang hiển thị?", _QMB.Yes|_QMB.No) != _QMB.Yes:
                return
            try:
                ids = []
                for i in range(table.rowCount()):
                    eid = table.item(i,0).data(Qt.UserRole) if table.item(i,0) else None
                    if eid:
                        ids.append(eid)
                if not ids:
                    return
                from .db import SessionLocal
                from .models import EmailLog
                s = SessionLocal();
                for eid in ids:
                    s.query(EmailLog).filter(EmailLog.id==eid).delete()
                s.commit(); s.close()
                load()
            except Exception as ex:
                QMessageBox.critical(dlg, "Lỗi", str(ex))
        btn_delete_all.clicked.connect(delete_all_filtered)
        def resend_all_filtered():
            try:
                total = 0; sent = 0; failed = 0
                if cb_all_scope.isChecked():
                    from .db import SessionLocal as _SL
                    from .models import EmailLog
                    s3 = _SL()
                    try:
                        q = s3.query(EmailLog)
                        t = type_box.currentText()
                        if not t.startswith("("):
                            q = q.filter(EmailLog.type == t)
                        u_sel = unit_box.currentText().strip() if unit_box.currentIndex() > -1 else ""
                        if u_sel and not u_sel.startswith("("):
                            q = q.filter(EmailLog.unit_name == u_sel)
                        else:
                            u = unit_edit.text().strip()
                            if u:
                                q = q.filter(EmailLog.unit_name.ilike(f"%{u}%"))
                        st = status_box.currentText()
                        if only_failed.isChecked():
                            q = q.filter(EmailLog.status == 'failed')
                        elif not st.startswith("("):
                            q = q.filter(EmailLog.status == st)
                        subj = subject_search.text().strip()
                        if subj:
                            q = q.filter(EmailLog.subject.ilike(f"%{subj}%"))
                        uidtxt = user_id_edit.text().strip()
                        if uidtxt.isdigit():
                            q = q.filter(EmailLog.user_id == int(uidtxt))
                        from datetime import datetime as _dt2
                        if e_from.isChecked():
                            fd = from_date.date(); q = q.filter(EmailLog.created_at >= _dt2(fd.year(), fd.month(), fd.day(), 0, 0, 0))
                        if e_to.isChecked():
                            td = to_date.date(); q = q.filter(EmailLog.created_at <= _dt2(td.year(), td.month(), td.day(), 23, 59, 59))
                        # Xác nhận số lượng
                        from PySide6.QtWidgets import QMessageBox as _QMB, QProgressDialog as _QPD
                        cnt = q.count()
                        if _QMB.question(dlg, "Xác nhận", f"Sẽ gửi lại {cnt} bản ghi phù hợp. Tiếp tục?", _QMB.Yes|_QMB.No) != _QMB.Yes:
                            return
                        total = cnt
                        pr = _QPD("Đang gửi lại...", "Hủy", 0, cnt, dlg); pr.setWindowModality(Qt.WindowModal)
                        rows = q.order_by(EmailLog.created_at.desc()).all()
                        idx = 0
                        for r in rows:
                            if pr.wasCanceled():
                                break
                            try:
                                t = getattr(r,'type','') or ''
                                unit_name = getattr(r,'unit_name','') or ''
                                subject = getattr(r,'subject','') or ''
                                body = getattr(r,'body','') or ''
                                paths = []
                                from pathlib import Path
                                for part in (getattr(r,'attachments','') or '').split(','):
                                    p = part.strip()
                                    if p:
                                        pp = Path(p)
                                        if pp.exists():
                                            paths.append(str(pp))
                                from .mailer import send_email_with_attachment, get_recipients_for_unit
                                recips = get_recipients_for_unit(unit_name) if t in ('salary_due','bhxh_monthly','contracts_expiring') and unit_name else None
                                ok = send_email_with_attachment(subject or '[HRMS] Resend', body or '', paths, to=recips)
                                # Log lại
                                try:
                                    s = _SL(); from .models import EmailLog as _EL
                                    s.add(_EL(type=t or 'generic', unit_name=(unit_name or None), recipients='', subject=subject or '[HRMS] Resend', body=(body or '')[:1000], attachments=', '.join(paths), status='sent' if ok else 'failed', user_id=self.current_user.get('id'))); s.commit(); s.close()
                                except Exception:
                                    pass
                                if ok: sent += 1
                                else: failed += 1
                            except Exception:
                                failed += 1
                            idx += 1; pr.setValue(idx); QApplication.processEvents()
                    finally:
                        try:
                            s3.close()
                        except Exception:
                            pass
                else:
                    from PySide6.QtWidgets import QProgressDialog as _QPD
                    n = table.rowCount()
                    pr2 = _QPD("Đang gửi lại...", "Hủy", 0, n, dlg); pr2.setWindowModality(Qt.WindowModal)
                    for i in range(n):
                        if pr2.wasCanceled():
                            break
                        total += 1
                        try:
                            t = table.item(i,1).text() if table.item(i,1) else ''
                            unit_name = table.item(i,2).text() if table.item(i,2) else ''
                            subject = table.item(i,3).text() if table.item(i,3) else ''
                            body = table.item(i,3).data(Qt.UserRole) if table.item(i,3) else ''
                            attach_disp = table.item(i,5).data(Qt.UserRole) if table.item(i,5) else (table.item(i,5).text() if table.item(i,5) else '')
                            paths = []
                            from pathlib import Path
                            for part in (attach_disp or '').split(','):
                                p = part.strip()
                                if p:
                                    pp = Path(p)
                                    if pp.exists():
                                        paths.append(str(pp))
                            recips = None
                            if t in ('salary_due','bhxh_monthly','contracts_expiring') and unit_name:
                                try:
                                    from .mailer import get_recipients_for_unit
                                    recips = get_recipients_for_unit(unit_name)
                                except Exception:
                                    recips = None
                            from .mailer import send_email_with_attachment
                            ok = send_email_with_attachment(subject or '[HRMS] Resend', body or '', paths, to=recips)
                            # Log lại
                            try:
                                from .db import SessionLocal
                                from .models import EmailLog
                                s = SessionLocal(); s.add(EmailLog(type=t or 'generic', unit_name=(unit_name or None), recipients='', subject=subject or '[HRMS] Resend', body=(body or '')[:1000], attachments=', '.join(paths), status='sent' if ok else 'failed', user_id=self.current_user.get('id'))); s.commit(); s.close()
                            except Exception:
                                pass
                            if ok: sent += 1
                            else: failed += 1
                        except Exception:
                            failed += 1
                        pr2.setValue(i+1); QApplication.processEvents()
                QMessageBox.information(dlg, "Kết quả", f"Tổng: {total}\nThành công: {sent}\nThất bại: {failed}")
            except Exception as ex:
                QMessageBox.critical(dlg, "Lỗi", str(ex))
        btn_resend_all.clicked.connect(resend_all_filtered)
        def resend_grouped_by_unit():
            try:
                # Gom theo đơn vị và loại hợp lệ
                allowed = {'salary_due','bhxh_monthly','contracts_expiring'}
                groups = {}
                from pathlib import Path
                if cb_all_scope.isChecked():
                    from .db import SessionLocal as _SL
                    from .models import EmailLog
                    s4 = _SL()
                    try:
                        q = s4.query(EmailLog)
                        t = type_box.currentText()
                        if not t.startswith("("):
                            q = q.filter(EmailLog.type == t)
                        u_sel = unit_box.currentText().strip() if unit_box.currentIndex() > -1 else ""
                        if u_sel and not u_sel.startswith("("):
                            q = q.filter(EmailLog.unit_name == u_sel)
                        else:
                            u = unit_edit.text().strip()
                            if u:
                                q = q.filter(EmailLog.unit_name.ilike(f"%{u}%"))
                        st = status_box.currentText()
                        if only_failed.isChecked():
                            q = q.filter(EmailLog.status == 'failed')
                        elif not st.startswith("("):
                            q = q.filter(EmailLog.status == st)
                        subj = subject_search.text().strip()
                        if subj:
                            q = q.filter(EmailLog.subject.ilike(f"%{subj}%"))
                        uidtxt = user_id_edit.text().strip()
                        if uidtxt.isdigit():
                            q = q.filter(EmailLog.user_id == int(uidtxt))
                        from datetime import datetime as _dt2
                        if e_from.isChecked():
                            fd = from_date.date(); q = q.filter(EmailLog.created_at >= _dt2(fd.year(), fd.month(), fd.day(), 0, 0, 0))
                        if e_to.isChecked():
                            td = to_date.date(); q = q.filter(EmailLog.created_at <= _dt2(td.year(), td.month(), td.day(), 23, 59, 59))
                        rows = q.order_by(EmailLog.created_at.desc()).all()
                        for r in rows:
                            t = getattr(r,'type','') or ''
                            if t not in allowed: continue
                            unit_name = getattr(r,'unit_name','') or ''
                            if not unit_name: continue
                            for part in (getattr(r,'attachments','') or '').split(','):
                                p = part.strip()
                                if p and Path(p).exists():
                                    key = (t, unit_name)
                                    groups.setdefault(key, set()).add(str(Path(p)))
                    finally:
                        try: s4.close()
                        except Exception: pass
                else:
                    for i in range(table.rowCount()):
                        t = table.item(i,1).text() if table.item(i,1) else ''
                        if t not in allowed: continue
                        unit_name = table.item(i,2).text() if table.item(i,2) else ''
                        if not unit_name: continue
                        attach_disp = table.item(i,5).data(Qt.UserRole) if table.item(i,5) else (table.item(i,5).text() if table.item(i,5) else '')
                        paths = []
                        for part in (attach_disp or '').split(','):
                            p = part.strip()
                            if p and Path(p).exists():
                                paths.append(str(Path(p)))
                        if not paths: continue
                        key = (t, unit_name)
                        groups.setdefault(key, set())
                        for p in paths:
                            groups[key].add(p)
                if not groups:
                    QMessageBox.information(dlg, "Không có dữ liệu", "Không có nhóm hợp lệ để gửi lại")
                    return
                # Thống kê và xác nhận
                def _hsz(n):
                    try:
                        for unit in ['B','KB','MB','GB','TB']:
                            if n < 1024.0:
                                return f"{n:0.1f} {unit}"
                            n /= 1024.0
                        return f"{n:0.1f} PB"
                    except Exception:
                        return str(n)
                total_groups = len(groups)
                total_files = 0
                total_bytes = 0
                from pathlib import Path as _P2
                for _files in groups.values():
                    total_files += len(_files)
                    for _f in _files:
                        try:
                            total_bytes += _P2(_f).stat().st_size
                        except Exception:
                            pass
                from PySide6.QtWidgets import QMessageBox as _QMB, QProgressDialog as _QPD
                if _QMB.question(dlg, "Xác nhận", f"Sẽ gửi lại {total_groups} nhóm, {total_files} tệp (~{_hsz(total_bytes)}). Tiếp tục?", _QMB.Yes|_QMB.No) != _QMB.Yes:
                    return
                sent = 0; failed = 0; total = total_groups
                from .mailer import get_recipients_for_unit, send_email_with_attachment, create_zip
                zip_when_group = QCheckBox.isChecked
                prg = _QPD("Đang gửi nhóm...", "Hủy", 0, total_groups, dlg); prg.setWindowModality(Qt.WindowModal)
                gi = 0
                for (t, unit_name), files in groups.items():
                    if prg.wasCanceled():
                        break
                    recips = get_recipients_for_unit(unit_name)
                    if not recips:
                        failed += 1; gi += 1; prg.setValue(gi); QApplication.processEvents(); continue
                    # subject with optional prefix/suffix
                    try:
                        subj_prefix = getattr(resend_grouped_by_unit, '__subj_prefix').text().strip() if hasattr(resend_grouped_by_unit, '__subj_prefix') else ''
                        subj_suffix = getattr(resend_grouped_by_unit, '__subj_suffix').text().strip() if hasattr(resend_grouped_by_unit, '__subj_suffix') else ''
                    except Exception:
                        subj_prefix = subj_suffix = ''
                    base = f"[HRMS] Resend {t} - {unit_name}"
                    tmpl = getattr(resend_grouped_by_unit, '__subj_template').text().strip() if hasattr(resend_grouped_by_unit, '__subj_template') else ''
                    if tmpl:
                        from datetime import date as _date
                        today = _date.today().isoformat()
                        subject = tmpl.replace('{type}', t).replace('{unit}', unit_name).replace('{count}', str(len(files))).replace('{date}', today)
                    else:
                        subject = f"{subj_prefix}{base}{subj_suffix}"
                    body = f"Gửi lại theo nhóm {t} cho {unit_name}. Đính kèm {len(files)} tệp."
                    file_list = sorted(list(files))
                    # ZIP tùy chọn nếu có nhiều tệp
                    ok = False
                    try:
                        if hasattr(resend_grouped_by_unit, '__zip_cb') and resend_grouped_by_unit.__zip_cb.isChecked() and len(file_list) > 1:
                            from pathlib import Path as _P
                            _P('exports').mkdir(exist_ok=True)
                            # tên ZIP tuỳ chọn
                            patt = resend_grouped_by_unit.__zip_pattern.text() if hasattr(resend_grouped_by_unit, '__zip_pattern') else ''
                            if not patt:
                                patt = "resend_{type}_{unit}_{ts}.zip"
                            from datetime import datetime as _dt
                            ts = _dt.now().strftime('%Y%m%d_%H%M%S')
                            safe_unit = unit_name.replace(' ', '_')
                            name = patt.replace('{type}', t).replace('{unit}', safe_unit).replace('{ts}', ts)
                            zip_path = _P('exports')/name
                            if create_zip(file_list, str(zip_path)):
                                ok = send_email_with_attachment(subject, body + "\n(Đính kèm ZIP)", [str(zip_path)], to=recips)
                            else:
                                ok = send_email_with_attachment(subject, body, file_list, to=recips)
                        else:
                            ok = send_email_with_attachment(subject, body, file_list, to=recips)
                    except Exception:
                        ok = False
                    # Log lại
                    try:
                        from .db import SessionLocal
                        from .models import EmailLog
                        s = SessionLocal(); s.add(EmailLog(type=t, unit_name=unit_name, recipients='', subject=subject, body=body, attachments=', '.join(sorted(list(files))), status='sent' if ok else 'failed', user_id=self.current_user.get('id'))); s.commit(); s.close()
                    except Exception:
                        pass
                    if ok: sent += 1
                    else: failed += 1
                    gi += 1; prg.setValue(gi); QApplication.processEvents()
                QMessageBox.information(dlg, "Kết quả", f"Nhóm: {total}\nThành công: {sent}\nThất bại: {failed}")
            except Exception as ex:
                QMessageBox.critical(dlg, "Lỗi", str(ex))
        # Thêm checkbox ZIP khi gửi nhóm
        from PySide6.QtWidgets import QCheckBox as _QCB
        _zip_cb = _QCB("ZIP khi gửi nhóm")
        resend_grouped_by_unit.__zip_cb = _zip_cb
        from PySide6.QtWidgets import QLineEdit as _QLE
        _zip_pattern = _QLE("resend_{type}_{unit}_{ts}.zip")
        resend_grouped_by_unit.__zip_pattern = _zip_pattern
        subj_prefix = _QLE(""); subj_suffix = _QLE(""); subj_template = _QLE("")
        resend_grouped_by_unit.__subj_prefix = subj_prefix
        resend_grouped_by_unit.__subj_suffix = subj_suffix
        resend_grouped_by_unit.__subj_template = subj_template
        f.addRow("", _zip_cb)
        f.addRow("Tên ZIP", _zip_pattern)
        f.addRow("Subject template", subj_template)
        f.addRow("Subject prefix", subj_prefix)
        f.addRow("Subject suffix", subj_suffix)
        btn_resend_group.clicked.connect(resend_grouped_by_unit)
        load()
        dlg.resize(1100, 600)
        dlg.exec()

    def manage_users(self):
        # Chỉ admin/hr mới được
        if (self.current_user.get('role') or '').lower() not in ('admin','hr'):
            QMessageBox.warning(self, "Không có quyền", "Bạn không có quyền quản lý người dùng")
            return
        # Dialog tạo nhanh người dùng
        dlg = QDialog(self)
        dlg.setWindowTitle("Thêm người dùng")
        f = QFormLayout(dlg)
        user_edit = QLineEdit()
        pw_edit = QLineEdit(); pw_edit.setEchoMode(QLineEdit.Password)
        role_box = QComboBox(); role_box.addItems(["user","hr","admin"])
        unit_box = QComboBox(); unit_box.addItem("-- Đơn vị --", None)
        # nạp đơn vị
        db = SessionLocal()
        from .models import Unit, User
        for u in db.query(Unit).order_by(Unit.name).all():
            unit_box.addItem(u.name, u.id)
        f.addRow("Username", user_edit)
        f.addRow("Password", pw_edit)
        f.addRow("Role", role_box)
        f.addRow("Đơn vị", unit_box)
        btn_ok = QPushButton("Tạo")
        btn_cancel = QPushButton("Hủy")
        row = QHBoxLayout(); row.addWidget(btn_ok); row.addWidget(btn_cancel)
        f.addRow(row)

        def do_create():
            username = user_edit.text().strip()
            pw = pw_edit.text()
            role = role_box.currentText()
            unit_id = unit_box.currentData()
            if not username or not pw:
                QMessageBox.warning(dlg, "Thiếu thông tin", "Nhập username và password")
                return
            # Tạo user
            from .security import hash_password
            if db.query(User).filter(User.username == username).first():
                QMessageBox.warning(dlg, "Tồn tại", "Username đã tồn tại")
                return
            u = User(username=username, password_hash=hash_password(pw), role=role, unit_id=unit_id)
            db.add(u); db.commit();
            try:
                log_action(db, self.current_user.get('id'), 'create_user', 'User', u.id, f"role={role};unit_id={unit_id}")
            except Exception:
                pass
            QMessageBox.information(dlg, "Thành công", "Đã tạo người dùng")
            dlg.accept()

        btn_ok.clicked.connect(do_create)
        btn_cancel.clicked.connect(dlg.reject)
        dlg.exec()
        db.close()

    def import_excel(self):
        from PySide6.QtWidgets import QFileDialog
        from .importer import import_persons_from_excel
        db = SessionLocal()
        try:
            xlsx_path, _ = QFileDialog.getOpenFileName(self, "Chọn file Excel", "", "Excel Files (*.xlsx)")
            if not xlsx_path:
                return
            # RBAC: chỉ admin/hr mới được import
            if (self.current_user.get('role') or '').lower() not in ('admin','hr'):
                QMessageBox.warning(self, "Không có quyền", "Bạn không có quyền import")
                return
            result = import_persons_from_excel(db, xlsx_path)
            if not result.get("ok"):
                QMessageBox.critical(self, "Import lỗi", result.get("error", "Lỗi không rõ"))
                return
            # Audit
            try:
                details = f"file={xlsx_path}; created={result['created']}; updated={result['updated']}"
                log_action(db, self.current_user.get('id'), 'import_persons', 'Person', None, details)
            except Exception:
                pass
            QMessageBox.information(self, "Import thành công", f"Tạo mới: {result['created']}, Cập nhật: {result['updated']}")
            self.refresh()
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", str(e))
        finally:
            if db:
                db.close()

    def _check_unit_permission(self, person) -> bool:
        role = (self.current_user.get('role') or '').lower()
        if role in ('admin','hr'):
            return True
        if role == 'unit_manager' and person and person.unit_id == self.current_user.get('unit_id'):
            return True
        return False

    def add_insurance(self):
        from PySide6.QtWidgets import QInputDialog
        from datetime import date
        from .insurance import add_insurance_event
        db, person = self.current_person()
        try:
            if not person:
                QMessageBox.information(self, "Chưa chọn", "Chọn một nhân sự trước")
                return
            etype, ok = QInputDialog.getText(self, "Loại sự kiện BHXH", "Nhập loại sự kiện:")
            if not ok or not etype:
                return
            if not self._check_unit_permission(person):
                QMessageBox.warning(self, "Không có quyền", "Chỉ quản lý đơn vị mới được thêm cho người thuộc đơn vị mình")
                return
            add_insurance_event(db, person, etype, date.today())
            # Audit
            try:
                log_action(db, self.current_user.get('id'), 'add_insurance_event', 'InsuranceEvent', None, f"type={etype}")
            except Exception:
                pass
            QMessageBox.information(self, "Thành công", "Đã thêm sự kiện BHXH")
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", str(e))
        finally:
            if db:
                db.close()

    def export_insurance(self):
        from datetime import date
        from .insurance import export_insurance_to_excel
        start = date(date.today().year, 1, 1)
        end = date(date.today().year, 12, 31)
        db = SessionLocal()
        try:
            Path("exports").mkdir(exist_ok=True)
            file_path = Path("exports") / f"bhxh_{start.year}.xlsx"
            export_insurance_to_excel(db, start, end, str(file_path), template_name=self.get_template_for('bhxh'), username=self.current_user.get('username'))
            # Audit
            try:
                log_action(db, self.current_user.get('id'), 'export_insurance', 'InsuranceEvent', None, f"file={file_path}")
            except Exception:
                pass
            QMessageBox.information(self, "Thành công", f"Đã xuất: {file_path}")
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", str(e))
        finally:
            db.close()

    def add_contract(self):
        from PySide6.QtWidgets import QInputDialog
        from datetime import date
        from .contracts import add_contract as _add_contract
        db, person = self.current_person()
        try:
            if not person:
                QMessageBox.information(self, "Chưa chọn", "Chọn một nhân sự trước")
                return
            if not self._check_unit_permission(person):
                QMessageBox.warning(self, "Không có quyền", "Chỉ quản lý đơn vị mới được thêm cho người thuộc đơn vị mình")
                return
            ctype, ok = QInputDialog.getText(self, "Loại HĐ", "Nhập loại hợp đồng:")
            if not ok or not ctype:
                return
            c = _add_contract(db, person, ctype, date.today())
            try:
                log_action(db, self.current_user.get('id'), 'add_contract', 'Contract', c.id, f"type={ctype}")
            except Exception:
                pass
            QMessageBox.information(self, "Thành công", "Đã thêm hợp đồng")
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", str(e))
        finally:
            if db:
                db.close()

    def export_contract(self):
        from .contracts import export_contracts_for_person
        db, person = self.current_person()
        try:
            if not person:
                QMessageBox.information(self, "Chưa chọn", "Chọn một nhân sự trước")
                return
            Path("exports").mkdir(exist_ok=True)
            file_path = Path("exports") / f"contracts_{person.code}.xlsx"
            export_contracts_for_person(db, person, str(file_path), template_name=self.get_template_for('contracts'), username=self.current_user.get('username'))
            try:
                log_action(db, self.current_user.get('id'), 'export_contracts', 'Person', person.id, f"file={file_path}")
            except Exception:
                pass
            QMessageBox.information(self, "Thành công", f"Đã xuất: {file_path}")
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", str(e))
        finally:
            if db:
                db.close()


def quarter_window(d):
    # cảnh báo vào ngày 15 của tháng 2/5/8/11, nhưng ở đây tạo cửa sổ quý
    q = (d.month - 1) // 3 + 1
    start_month = 3 * (q - 1) + 1
    from datetime import date as _date
    start = _date(d.year, start_month, 1)
    # end: ngày cuối tháng trong quý
    end_month = start_month + 2
    from calendar import monthrange
    last_day = monthrange(d.year, end_month)[1]
    end = _date(d.year, end_month, last_day)
    return start, end


def main():
    # Khởi tạo DB nếu chưa có và seed dữ liệu demo lần đầu
    init_db()
    seed_basic_data()

    # Khởi động scheduler cảnh báo in-app
    try:
        from .scheduler import schedule_jobs
        schedule_jobs()
    except Exception as e:
        print(f"[WARN] Scheduler not started: {e}")

    app = QApplication(sys.argv)
    w = LoginWindow()
    w.show()
    sys.exit(app.exec())
