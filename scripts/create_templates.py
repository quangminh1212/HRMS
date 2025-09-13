from docx import Document
from pathlib import Path

def create_profile_template():
    Path('templates').mkdir(exist_ok=True)
    p = Document()
    p.add_heading('TRÍCH NGANG NHÂN SỰ', level=1)
    rows = [
        ('Họ và tên', '{{full_name}}'),
        ('Mã', '{{code}}'),
        ('Ngày sinh', '{{dob}}'),
        ('Giới tính', '{{gender}}'),
        ('Dân tộc', '{{ethnicity}}'),
        ('Tôn giáo', '{{religion}}'),
        ('Quê quán', '{{hometown}}'),
        ('Chức danh', '{{position}}'),
        ('Đơn vị', '{{unit}}'),
        ('Ngày vào Đảng', '{{party_joined_date}}'),
        ('LLCT', '{{llct_level}}'),
        ('Trình độ chuyên môn', '{{professional_level}}'),
        ('Tình trạng công tác', '{{status}}'),
        ('Điện thoại', '{{phone}}'),
        ('Email', '{{email}}'),
    ]
    t = p.add_table(rows=1, cols=2)
    h = t.rows[0].cells
    h[0].text = 'Trường'
    h[1].text = 'Giá trị'
    for k, v in rows:
        row = t.add_row().cells
        row[0].text = k
        row[1].text = v
    Path('templates').mkdir(exist_ok=True)
    p.save('templates/trich_ngang_template.docx')


def create_letter_templates():
    Path('templates').mkdir(exist_ok=True)
    # Minimal templates with placeholders
    letters = {
        'cong_van_ra_soat_nang_luong.docx': [
            'CÔNG VĂN RÀ SOÁT NÂNG LƯƠNG',
            'Kính gửi: {{org_name}}',
            'V/v rà soát nâng lương định kỳ quý {{quarter}} năm {{year}}',
            'Danh sách kèm theo file Excel.',
            'Ngày {{date}}'
        ],
        'thong_bao_nang_luong.docx': [
            'THÔNG BÁO NÂNG LƯƠNG',
            'Ông/Bà: {{full_name}} ({{code}})',
            'Đơn vị: {{unit}}',
            'Từ ngày: {{effective_date}}',
            'Bậc/hệ số mới: {{step}}/{{coefficient}}'
        ],
        'quyet_dinh_nang_luong.docx': [
            'QUYẾT ĐỊNH NÂNG LƯƠNG',
            'Căn cứ: ...',
            'Quyết định nâng lương cho Ông/Bà {{full_name}} từ ngày {{effective_date}}.',
            'Bậc/hệ số mới: {{step}}/{{coefficient}}'
        ],
        'thong_bao_nghi_huu.docx': [
            'THÔNG BÁO NGHỈ HƯU',
            'Ông/Bà: {{full_name}} ({{code}})',
            'Ngày nghỉ hưu dự kiến: {{retirement_date}}',
            'Ngày thông báo: {{date}}'
        ],
        'quyet_dinh_nghi_huu.docx': [
            'QUYẾT ĐỊNH NGHỈ HƯU',
            'Căn cứ: ...',
            'Quyết định cho Ông/Bà {{full_name}} nghỉ hưu kể từ ngày {{retirement_date}}.'
        ],
    }
    for fname, lines in letters.items():
        d = Document()
        for line in lines:
            d.add_paragraph(line)
        d.save(str(Path('templates') / fname))


if __name__ == '__main__':
    create_profile_template()
    create_letter_templates()
