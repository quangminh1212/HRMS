"""
Các hàm tiện ích cho HRMS Streamlit Version
Bao gồm tất cả logic nghiệp vụ theo yêu cầu
"""

from datetime import datetime, timedelta, date
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
import io
import os

def calculate_retirement_date(birth_date, gender):
    """
    Tính ngày nghỉ hưu theo quy định mới
    Nam: 60 tuổi 3 tháng (tăng dần từ 2021)
    Nữ: 55 tuổi 4 tháng (tăng dần từ 2021)
    """
    if not birth_date:
        return None
        
    if gender == 'Nam':
        retirement_age_years = 60
        retirement_age_months = 3
    else:
        retirement_age_years = 55
        retirement_age_months = 4
    
    # Tính ngày nghỉ hưu
    retirement_date = birth_date.replace(year=birth_date.year + retirement_age_years)
    
    # Cộng thêm số tháng
    if retirement_date.month + retirement_age_months <= 12:
        retirement_date = retirement_date.replace(month=retirement_date.month + retirement_age_months)
    else:
        retirement_date = retirement_date.replace(
            year=retirement_date.year + 1,
            month=retirement_date.month + retirement_age_months - 12
        )
    
    return retirement_date

def check_salary_increase_eligibility(employee):
    """
    Kiểm tra điều kiện nâng lương thường xuyên
    - Chuyên viên trở lên: 36 tháng
    - Nhân viên, Thủ quỹ: 24 tháng
    """
    if not employee.current_salary_date:
        # Nếu chưa có lần nâng lương nào, tính từ ngày bắt đầu công tác
        reference_date = employee.organization_start_date
    else:
        reference_date = employee.current_salary_date
    
    if not reference_date:
        return False
    
    # Tính số tháng từ lần nâng lương gần nhất
    months_since_last = calculate_months_difference(reference_date, date.today())
    
    # Xác định thời gian cần thiết dựa trên chức vụ
    if employee.position and any(pos in employee.position 
                                for pos in ['Chuyên viên', 'Chuyên gia', 'Giám đốc', 'Phó giám đốc', 'Trưởng phòng']):
        required_months = 36
    else:
        required_months = 24
    
    return months_since_last >= required_months, months_since_last

def calculate_months_difference(start_date, end_date):
    """Tính số tháng giữa 2 ngày"""
    return (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month)

def calculate_seniority_allowance(employee):
    """
    Tính phụ cấp thâm niên vượt khung
    - Hưởng 5% khi đủ 36 tháng ở bậc lương cuối cùng (Chuyên viên) hoặc 24 tháng (Nhân viên)
    - Sau đó mỗi năm tăng 1%
    """
    if not employee.current_salary_date:
        return 0
    
    # Kiểm tra xem có phải bậc lương cuối cùng không (giả sử hệ số >= 4.0 là bậc cuối)
    if (employee.current_salary_coefficient or 0) < 4.0:
        return 0
    
    months_at_max_level = calculate_months_difference(employee.current_salary_date, date.today())
    
    # Xác định thời gian cần thiết
    if employee.position and 'Chuyên viên' in employee.position:
        required_months = 36
    else:
        required_months = 24
    
    if months_at_max_level < required_months:
        return 0
    
    # Tính phụ cấp: 5% ban đầu + 1% cho mỗi năm sau đó
    years_beyond_requirement = (months_at_max_level - required_months) // 12
    seniority_allowance = 5 + years_beyond_requirement
    
    return min(seniority_allowance, 30)  # Giới hạn tối đa 30%

def check_appointment_eligibility(employee, target_position):
    """
    Kiểm tra điều kiện bổ nhiệm
    - Trong quy hoạch
    - Đáp ứng yêu cầu văn bằng chứng chỉ
    - Kinh nghiệm công tác
    - Độ tuổi phù hợp
    """
    eligibility = {
        'is_eligible': False,
        'in_planning': False,
        'education_met': False,
        'experience_met': False,
        'age_appropriate': False,
        'messages': []
    }
    
    # Kiểm tra quy hoạch
    if employee.current_planning_position and target_position in employee.current_planning_position:
        eligibility['in_planning'] = True
    else:
        eligibility['messages'].append(f"Không có trong quy hoạch cho vị trí {target_position}")
    
    # Kiểm tra trình độ học vấn
    if employee.professional_level:
        if 'Đại học' in employee.professional_level or 'Thạc sĩ' in employee.professional_level:
            eligibility['education_met'] = True
        else:
            eligibility['messages'].append("Chưa đạt trình độ học vấn yêu cầu")
    else:
        eligibility['messages'].append("Chưa có thông tin trình độ học vấn")
    
    # Kiểm tra kinh nghiệm
    if employee.start_date:
        years_experience = (date.today() - employee.start_date).days / 365
        if years_experience >= 5:  # Giả sử cần ít nhất 5 năm kinh nghiệm
            eligibility['experience_met'] = True
        else:
            eligibility['messages'].append(f"Kinh nghiệm chưa đủ: {years_experience:.1f} năm (cần ít nhất 5 năm)")
    
    # Kiểm tra tuổi
    if employee.date_of_birth:
        current_age = (date.today() - employee.date_of_birth).days / 365
        if 30 <= current_age <= 50:  # Giả sử độ tuổi phù hợp cho bổ nhiệm
            eligibility['age_appropriate'] = True
        else:
            eligibility['messages'].append(f"Độ tuổi không phù hợp: {current_age:.0f} tuổi")
    
    # Tổng hợp kết quả
    eligibility['is_eligible'] = all([
        eligibility['in_planning'],
        eligibility['education_met'],
        eligibility['experience_met'],
        eligibility['age_appropriate']
    ])
    
    return eligibility

def get_salary_increase_schedule():
    """
    Lấy lịch cảnh báo nâng lương
    Cảnh báo vào ngày 15 của tháng 2, 5, 8, 11 (trước đợt xét nâng lương cuối quý)
    """
    current_year = date.today().year
    schedule = []
    
    alert_months = [2, 5, 8, 11]  # Tháng 2, 5, 8, 11
    review_months = [3, 6, 9, 12]  # Tháng xét nâng lương
    
    for i, alert_month in enumerate(alert_months):
        alert_date = date(current_year, alert_month, 15)
        review_date = date(current_year, review_months[i], 28)  # Cuối tháng xét
        
        schedule.append({
            'alert_date': alert_date,
            'review_date': review_date,
            'quarter': f'Q{(review_months[i]-1)//3 + 1}',
            'description': f'Cảnh báo cho đợt xét nâng lương Q{(review_months[i]-1)//3 + 1}/{current_year}'
        })
    
    return schedule

def check_retirement_notifications_needed(employees):
    """
    Kiểm tra nhân viên cần thông báo/quyết định nghỉ hưu
    - Thông báo: trước 6 tháng
    - Quyết định: trước 3 tháng
    """
    today = date.today()
    notification_6_months = today + timedelta(days=180)  # 6 tháng
    decision_3_months = today + timedelta(days=90)      # 3 tháng
    
    results = {
        'need_notification': [],
        'need_decision': [],
        'need_early_salary_review': []  # Cần xem xét nâng lương trước thời hạn
    }
    
    for emp in employees:
        if emp.retirement_date and emp.status == 'active':
            if emp.retirement_date <= notification_6_months:
                results['need_notification'].append(emp)
            
            if emp.retirement_date <= decision_3_months:
                results['need_decision'].append(emp)
                
                # Kiểm tra nâng lương trước thời hạn cho người sắp nghỉ hưu
                if check_salary_increase_eligibility(emp):
                    results['need_early_salary_review'].append(emp)
    
    return results

def export_employee_word(employee):
    """
    Xuất thông tin nhân sự ra file Word theo mẫu
    """
    doc = Document()
    
    # Tiêu đề
    title = doc.add_heading('THÔNG TIN NHÂN SỰ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Thông tin cá nhân
    doc.add_heading('I. THÔNG TIN CÁ NHÂN', level=1)
    
    info_table = doc.add_table(rows=7, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    # Điền thông tin vào bảng
    cells = info_table.rows
    cells[0].cells[0].text = 'Mã nhân viên:'
    cells[0].cells[1].text = employee.employee_code or ''
    cells[1].cells[0].text = 'Họ và tên:'
    cells[1].cells[1].text = employee.full_name or ''
    cells[2].cells[0].text = 'Ngày sinh:'
    cells[2].cells[1].text = employee.date_of_birth.strftime('%d/%m/%Y') if employee.date_of_birth else ''
    cells[3].cells[0].text = 'Giới tính:'
    cells[3].cells[1].text = employee.gender or ''
    cells[4].cells[0].text = 'Dân tộc:'
    cells[4].cells[1].text = employee.ethnicity or ''
    cells[5].cells[0].text = 'Tôn giáo:'
    cells[5].cells[1].text = employee.religion or ''
    cells[6].cells[0].text = 'Quê quán:'
    cells[6].cells[1].text = employee.hometown or ''
    
    # Thông tin công việc
    doc.add_heading('II. THÔNG TIN CÔNG VIỆC', level=1)
    
    work_table = doc.add_table(rows=6, cols=2)
    work_table.style = 'Light Grid Accent 1'
    
    work_cells = work_table.rows
    work_cells[0].cells[0].text = 'Chức vụ:'
    work_cells[0].cells[1].text = employee.position or ''
    work_cells[1].cells[0].text = 'Đơn vị:'
    work_cells[1].cells[1].text = employee.department or ''
    work_cells[2].cells[0].text = 'Ngày vào Đảng:'
    work_cells[2].cells[1].text = employee.party_join_date.strftime('%d/%m/%Y') if employee.party_join_date else ''
    work_cells[3].cells[0].text = 'Trình độ LLCT:'
    work_cells[3].cells[1].text = employee.political_theory_level or ''
    work_cells[4].cells[0].text = 'Trình độ chuyên môn:'
    work_cells[4].cells[1].text = employee.professional_level or ''
    work_cells[5].cells[0].text = 'Ngày nghỉ hưu dự kiến:'
    work_cells[5].cells[1].text = employee.retirement_date.strftime('%d/%m/%Y') if employee.retirement_date else ''
    
    # Thông tin lương
    doc.add_heading('III. THÔNG TIN LƯƠNG', level=1)
    
    salary_table = doc.add_table(rows=4, cols=2)
    salary_table.style = 'Light Grid Accent 1'
    
    salary_cells = salary_table.rows
    salary_cells[0].cells[0].text = 'Ngạch lương:'
    salary_cells[0].cells[1].text = employee.current_salary_level or ''
    salary_cells[1].cells[0].text = 'Hệ số lương:'
    salary_cells[1].cells[1].text = str(employee.current_salary_coefficient) if employee.current_salary_coefficient else ''
    salary_cells[2].cells[0].text = 'Phụ cấp chức vụ:'
    salary_cells[2].cells[1].text = str(employee.position_allowance) if employee.position_allowance else '0'
    salary_cells[3].cells[0].text = 'Ngày nâng lương gần nhất:'
    salary_cells[3].cells[1].text = employee.last_salary_increase_date.strftime('%d/%m/%Y') if employee.last_salary_increase_date else ''
    
    # Thông tin liên hệ
    doc.add_heading('IV. THÔNG TIN LIÊN HỆ', level=1)
    
    contact_table = doc.add_table(rows=3, cols=2)
    contact_table.style = 'Light Grid Accent 1'
    
    contact_cells = contact_table.rows
    contact_cells[0].cells[0].text = 'Điện thoại:'
    contact_cells[0].cells[1].text = employee.phone or ''
    contact_cells[1].cells[0].text = 'Email:'
    contact_cells[1].cells[1].text = employee.email or ''
    contact_cells[2].cells[0].text = 'Địa chỉ:'
    contact_cells[2].cells[1].text = employee.address or ''
    
    # Lưu file vào memory
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    
    return bio

def export_salary_decision(employees_list):
    """
    Xuất quyết định nâng lương ra file Word
    """
    doc = Document()
    
    # Header
    header = doc.add_paragraph()
    header.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n').bold = True
    header.add_run('Độc lập - Tự do - Hạnh phúc\n').italic = True
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('_' * 30).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Title
    title = doc.add_heading('QUYẾT ĐỊNH', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Về việc nâng bậc lương thường xuyên')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    
    # Số quyết định
    doc.add_paragraph(f'Số: {generate_decision_number()}/QĐ-{datetime.now().year}')
    
    # Căn cứ pháp lý
    doc.add_heading('GIÁM ĐỐC', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Căn cứ Nghị định số 204/2004/NĐ-CP ngày 14/12/2004 của Chính phủ về chế độ tiền lương;')
    doc.add_paragraph('Căn cứ Thông tư số 08/2013/TT-BNV ngày 31/7/2013 của Bộ Nội vụ;')
    doc.add_paragraph('Xét đề nghị của Trưởng phòng Tổ chức - Hành chính,')
    
    doc.add_heading('QUYẾT ĐỊNH:', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Điều 1. Nâng bậc lương thường xuyên cho các ông (bà) có tên sau:')
    
    # Tạo bảng danh sách
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'
    
    # Header bảng
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'STT'
    hdr_cells[1].text = 'Họ và tên'
    hdr_cells[2].text = 'Chức vụ'
    hdr_cells[3].text = 'Bậc lương cũ'
    hdr_cells[4].text = 'Bậc lương mới'
    hdr_cells[5].text = 'Thời điểm hưởng'
    
    # Thêm dữ liệu
    for idx, emp in enumerate(employees_list, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = emp['name']
        row_cells[2].text = emp['position']
        row_cells[3].text = str(emp['old_coefficient'])
        row_cells[4].text = str(emp['new_coefficient'])
        row_cells[5].text = emp['effective_date'].strftime('%d/%m/%Y')
    
    doc.add_paragraph()
    doc.add_paragraph('Điều 2. Các ông (bà) có tên tại Điều 1 được hưởng lương mới kể từ ngày có hiệu lực của Quyết định này.')
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run(f'..., ngày {datetime.now().day} tháng {datetime.now().month} năm {datetime.now().year}').italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    signature = doc.add_paragraph('GIÁM ĐỐC')
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signature.runs[0].bold = True
    
    # Lưu vào memory
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    
    return bio

def export_insurance_excel(employees_list, changes_type):
    """
    Xuất file Excel để báo bảo hiểm
    changes_type: 'salary_change', 'retirement', 'maternity', 'sick_leave', 'study_leave'
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"Báo BHXH - {changes_type}"
    
    # Headers tùy theo loại thay đổi
    if changes_type == 'salary_change':
        headers = [
            'STT', 'Mã NV', 'Họ tên', 'Số sổ BHXH', 
            'Lương cũ', 'Lương mới', 'Từ tháng/năm', 'Ghi chú'
        ]
    elif changes_type == 'retirement':
        headers = [
            'STT', 'Mã NV', 'Họ tên', 'Số sổ BHXH',
            'Ngày nghỉ hưu', 'Lương cuối', 'Ghi chú'
        ]
    elif changes_type == 'maternity':
        headers = [
            'STT', 'Mã NV', 'Họ tên', 'Số sổ BHXH',
            'Từ ngày', 'Đến ngày', 'Chế độ', 'Ghi chú'
        ]
    
    # Định dạng header
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Thêm dữ liệu
    for idx, emp in enumerate(employees_list, 2):
        ws.cell(row=idx, column=1, value=idx-1)
        ws.cell(row=idx, column=2, value=emp.get('employee_code'))
        ws.cell(row=idx, column=3, value=emp.get('full_name'))
        ws.cell(row=idx, column=4, value=emp.get('insurance_number'))
        
        if changes_type == 'salary_change':
            ws.cell(row=idx, column=5, value=emp.get('old_salary'))
            ws.cell(row=idx, column=6, value=emp.get('new_salary'))
            ws.cell(row=idx, column=7, value=emp.get('effective_month'))
            ws.cell(row=idx, column=8, value=emp.get('notes'))
    
    # Điều chỉnh độ rộng cột
    for column_cells in ws.columns:
        length = max(len(str(cell.value or '')) for cell in column_cells)
        ws.column_dimensions[column_cells[0].column_letter].width = min(length + 2, 50)
    
    # Lưu vào memory
    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)
    
    return bio

def generate_decision_number():
    """Tạo số quyết định tự động"""
    return f"{datetime.now().strftime('%m%d')}{datetime.now().strftime('%H%M')}"

def calculate_work_experience(start_date, end_date=None):
    """Tính thời gian công tác"""
    if not start_date:
        return 0
    
    if not end_date:
        end_date = date.today()
    
    return (end_date - start_date).days / 365

def check_planning_quota(position, department, current_count):
    """
    Kiểm tra số lượng quy hoạch cho một vị trí có vượt quá không
    """
    # Định nghĩa quota cho từng vị trí (có thể lưu trong config hoặc database)
    position_quotas = {
        'Trưởng phòng': 1,
        'Phó trưởng phòng': 2,
        'Chuyên viên cao cấp': 3,
        'Chuyên viên chính': 5,
        'Chuyên viên': 10
    }
    
    department_multiplier = {
        'Phòng Tổ chức - Hành chính': 1.2,
        'Phòng Tài chính - Kế toán': 1.0,
        'Phòng Kinh doanh': 1.5
    }
    
    base_quota = position_quotas.get(position, 1)
    multiplier = department_multiplier.get(department, 1.0)
    max_quota = int(base_quota * multiplier)
    
    return {
        'current_count': current_count,
        'max_quota': max_quota,
        'is_exceeded': current_count > max_quota,
        'available_slots': max(0, max_quota - current_count)
    }
