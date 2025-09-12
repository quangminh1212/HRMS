"""
Các hàm tiện ích cho HRMS
"""

from datetime import datetime, timedelta, date
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
import os

def calculate_retirement_date(birth_date, gender):
    """
    Tính ngày nghỉ hưu dựa trên ngày sinh và giới tính
    Nam: 60 tuổi + 3 tháng (từ 2021)
    Nữ: 55 tuổi + 4 tháng (từ 2021) 
    """
    if gender == 'Nam':
        retirement_age_years = 60
        retirement_age_months = 3
    else:
        retirement_age_years = 55
        retirement_age_months = 4
    
    retirement_date = birth_date.replace(year=birth_date.year + retirement_age_years)
    retirement_date = retirement_date + timedelta(days=retirement_age_months * 30)
    
    return retirement_date

def check_salary_increase_eligibility(employee):
    """
    Kiểm tra điều kiện nâng lương
    - Chuyên viên trở lên: 36 tháng
    - Nhân viên: 24 tháng
    """
    if not employee.last_salary_increase_date:
        # Nếu chưa có lần nâng lương nào, tính từ ngày bắt đầu
        last_date = employee.start_date
    else:
        last_date = employee.last_salary_increase_date
    
    if not last_date:
        return False
    
    months_since_last = (datetime.now().date() - last_date).days / 30
    
    # Xác định thời gian cần thiết dựa trên ngạch
    if employee.current_salary_level and 'Chuyên viên' in employee.current_salary_level:
        required_months = 36
    else:
        required_months = 24
    
    return months_since_last >= required_months

def export_to_word(employee):
    """
    Xuất thông tin nhân sự ra file Word
    """
    doc = Document()
    
    # Tiêu đề
    title = doc.add_heading('THÔNG TIN NHÂN SỰ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Thông tin cá nhân
    doc.add_heading('I. THÔNG TIN CÁ NHÂN', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Họ và tên: ').bold = True
    p.add_run(employee.full_name)
    
    p = doc.add_paragraph()
    p.add_run('Mã nhân viên: ').bold = True
    p.add_run(employee.employee_code)
    
    p = doc.add_paragraph()
    p.add_run('Ngày sinh: ').bold = True
    p.add_run(employee.date_of_birth.strftime('%d/%m/%Y') if employee.date_of_birth else '')
    
    p = doc.add_paragraph()
    p.add_run('Giới tính: ').bold = True
    p.add_run(employee.gender or '')
    
    p = doc.add_paragraph()
    p.add_run('Dân tộc: ').bold = True
    p.add_run(employee.ethnicity or '')
    
    p = doc.add_paragraph()
    p.add_run('Tôn giáo: ').bold = True
    p.add_run(employee.religion or '')
    
    p = doc.add_paragraph()
    p.add_run('Quê quán: ').bold = True
    p.add_run(employee.hometown or '')
    
    # Thông tin công việc
    doc.add_heading('II. THÔNG TIN CÔNG VIỆC', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Chức vụ: ').bold = True
    p.add_run(employee.position or '')
    
    p = doc.add_paragraph()
    p.add_run('Đơn vị: ').bold = True
    p.add_run(employee.department or '')
    
    p = doc.add_paragraph()
    p.add_run('Ngày vào Đảng: ').bold = True
    p.add_run(employee.party_join_date.strftime('%d/%m/%Y') if employee.party_join_date else '')
    
    p = doc.add_paragraph()
    p.add_run('Trình độ lý luận chính trị: ').bold = True
    p.add_run(employee.political_theory_level or '')
    
    p = doc.add_paragraph()
    p.add_run('Trình độ chuyên môn: ').bold = True
    p.add_run(employee.professional_level or '')
    
    # Thông tin lương
    doc.add_heading('III. THÔNG TIN LƯƠNG', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Ngạch lương: ').bold = True
    p.add_run(employee.current_salary_level or '')
    
    p = doc.add_paragraph()
    p.add_run('Hệ số lương: ').bold = True
    p.add_run(str(employee.current_salary_coefficient) if employee.current_salary_coefficient else '')
    
    p = doc.add_paragraph()
    p.add_run('Ngày nâng lương gần nhất: ').bold = True
    p.add_run(employee.last_salary_increase_date.strftime('%d/%m/%Y') if employee.last_salary_increase_date else '')
    
    # Thông tin liên hệ
    doc.add_heading('IV. THÔNG TIN LIÊN HỆ', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Điện thoại: ').bold = True
    p.add_run(employee.phone or '')
    
    p = doc.add_paragraph()
    p.add_run('Email: ').bold = True
    p.add_run(employee.email or '')
    
    # Lưu file
    if not os.path.exists('exports'):
        os.makedirs('exports')
    
    filename = f'exports/ThongTinNhanSu_{employee.employee_code}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    doc.save(filename)
    
    return filename

def export_to_excel(employees):
    """
    Xuất danh sách nhân viên ra Excel
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Danh sách nhân sự"
    
    # Headers
    headers = [
        'STT', 'Mã NV', 'Họ và tên', 'Ngày sinh', 'Giới tính',
        'Chức vụ', 'Đơn vị', 'Ngạch lương', 'Hệ số lương',
        'Điện thoại', 'Email', 'Trạng thái'
    ]
    
    # Định dạng header
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Thêm dữ liệu
    for idx, emp in enumerate(employees, 2):
        ws.cell(row=idx, column=1, value=idx-1)
        ws.cell(row=idx, column=2, value=emp.employee_code)
        ws.cell(row=idx, column=3, value=emp.full_name)
        ws.cell(row=idx, column=4, value=emp.date_of_birth.strftime('%d/%m/%Y') if emp.date_of_birth else '')
        ws.cell(row=idx, column=5, value=emp.gender)
        ws.cell(row=idx, column=6, value=emp.position)
        ws.cell(row=idx, column=7, value=emp.department)
        ws.cell(row=idx, column=8, value=emp.current_salary_level)
        ws.cell(row=idx, column=9, value=emp.current_salary_coefficient)
        ws.cell(row=idx, column=10, value=emp.phone)
        ws.cell(row=idx, column=11, value=emp.email)
        ws.cell(row=idx, column=12, value=emp.status)
    
    # Điều chỉnh độ rộng cột
    for column_cells in ws.columns:
        length = max(len(str(cell.value or '')) for cell in column_cells)
        ws.column_dimensions[column_cells[0].column_letter].width = min(length + 2, 50)
    
    # Lưu file
    if not os.path.exists('exports'):
        os.makedirs('exports')
    
    filename = f'exports/DanhSachNhanSu_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
    wb.save(filename)
    
    return filename

def create_salary_decision_document(employees):
    """
    Tạo quyết định nâng lương
    """
    doc = Document()
    
    # Header
    header = doc.add_paragraph()
    header.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n').bold = True
    header.add_run('Độc lập - Tự do - Hạnh phúc\n').italic = True
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('_' * 30).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Title
    title = doc.add_heading('QUYẾT ĐỊNH', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Về việc nâng bậc lương thường xuyên')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].italic = True
    
    # Content
    doc.add_heading('GIÁM ĐỐC', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Căn cứ Nghị định số 204/2004/NĐ-CP ngày 14/12/2004 của Chính phủ về chế độ tiền lương;')
    doc.add_paragraph('Căn cứ Thông tư số 08/2013/TT-BNV ngày 31/7/2013 của Bộ Nội vụ;')
    doc.add_paragraph('Xét đề nghị của Trưởng phòng Tổ chức - Hành chính,')
    
    doc.add_heading('QUYẾT ĐỊNH:', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Điều 1. Nâng bậc lương thường xuyên cho các ông (bà) có tên sau:')
    
    # Tạo bảng
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'
    
    # Header bảng
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'STT'
    hdr_cells[1].text = 'Họ và tên'
    hdr_cells[2].text = 'Chức vụ'
    hdr_cells[3].text = 'Bậc lương cũ'
    hdr_cells[4].text = 'Bậc lương mới'
    hdr_cells[5].text = 'Thời điểm hưởng'
    
    # Thêm dữ liệu nhân viên
    for idx, emp in enumerate(employees, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = emp['name']
        row_cells[2].text = emp['position']
        row_cells[3].text = emp['old_salary']
        row_cells[4].text = emp['new_salary']
        row_cells[5].text = emp['effective_date']
    
    doc.add_paragraph()
    doc.add_paragraph('Điều 2. Các ông (bà) có tên tại Điều 1 được hưởng lương mới kể từ ngày có hiệu lực của Quyết định này.')
    doc.add_paragraph('Điều 3. Các ông (bà) Trưởng phòng Tổ chức - Hành chính, Trưởng phòng Tài chính - Kế toán và các ông (bà) có tên tại Điều 1 chịu trách nhiệm thi hành Quyết định này.')
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run(f'..., ngày {datetime.now().day} tháng {datetime.now().month} năm {datetime.now().year}').italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    signature = doc.add_paragraph('GIÁM ĐỐC')
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signature.runs[0].bold = True
    
    # Lưu file
    if not os.path.exists('exports'):
        os.makedirs('exports')
    
    filename = f'exports/QuyetDinhNangLuong_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    doc.save(filename)
    
    return filename

def calculate_insurance_contribution(salary, insurance_type='BHXH'):
    """
    Tính đóng bảo hiểm
    BHXH: 8% người lao động, 17.5% người sử dụng lao động
    BHYT: 1.5% người lao động, 3% người sử dụng lao động  
    BHTN: 1% người lao động, 1% người sử dụng lao động
    """
    rates = {
        'BHXH': {'employee': 0.08, 'employer': 0.175},
        'BHYT': {'employee': 0.015, 'employer': 0.03},
        'BHTN': {'employee': 0.01, 'employer': 0.01}
    }
    
    if insurance_type in rates:
        employee_contribution = salary * rates[insurance_type]['employee']
        employer_contribution = salary * rates[insurance_type]['employer']
        return {
            'employee': employee_contribution,
            'employer': employer_contribution,
            'total': employee_contribution + employer_contribution
        }
    
    return None

def check_planning_eligibility(employee, position):
    """
    Kiểm tra điều kiện quy hoạch
    """
    eligibility = {
        'age_valid': True,
        'education_valid': True,
        'experience_valid': True,
        'political_valid': True,
        'messages': []
    }
    
    # Kiểm tra tuổi
    if employee.date_of_birth:
        age = (datetime.now().date() - employee.date_of_birth).days / 365
        if age > 55:  # Giả sử tuổi quy hoạch tối đa là 55
            eligibility['age_valid'] = False
            eligibility['messages'].append(f'Tuổi hiện tại ({int(age)}) vượt quá tuổi quy hoạch')
    
    # Kiểm tra trình độ
    if not employee.professional_level or 'Đại học' not in employee.professional_level:
        eligibility['education_valid'] = False
        eligibility['messages'].append('Chưa có bằng đại học')
    
    # Kiểm tra kinh nghiệm
    if employee.start_date:
        years_experience = (datetime.now().date() - employee.start_date).days / 365
        if years_experience < 3:  # Giả sử cần ít nhất 3 năm kinh nghiệm
            eligibility['experience_valid'] = False
            eligibility['messages'].append(f'Kinh nghiệm chưa đủ (cần ít nhất 3 năm)')
    
    # Kiểm tra lý luận chính trị
    if 'Trưởng phòng' in position and not employee.political_theory_level:
        eligibility['political_valid'] = False
        eligibility['messages'].append('Chưa có chứng chỉ lý luận chính trị')
    
    eligibility['is_eligible'] = all([
        eligibility['age_valid'],
        eligibility['education_valid'],
        eligibility['experience_valid'],
        eligibility['political_valid']
    ])
    
    return eligibility
